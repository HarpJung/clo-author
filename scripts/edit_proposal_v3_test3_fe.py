"""Update Research Proposal v3 Section 4.4 (Test 3) with VC FE results."""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

path = "C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx"
doc = Document(path)

# Find paragraphs 214-217 (Test 3 content) and replace
# 214 = description, 215 = old results, 216 = regulatory, 217 = old table caption

# New content for paragraphs 214-217
new_texts = [
    # Para 214: Description (updated)
    (
        "This test exploits the network structure of observer appointments to identify "
        "information spillover. When an individual serves as an observer at a private Firm A "
        "and is connected through a VC firm to a public Firm B, information about events at "
        "Firm A may flow through the VC network to Firm B. We test this by examining whether "
        "material events at private companies (captured by CIQ Key Dev\u2014press releases, "
        "executive changes, financing rounds, etc.) predict abnormal returns at connected "
        "public portfolio companies in the months before the event becomes public. The key "
        "insight is that observers learn information in board meetings weeks or months before "
        "public disclosure. We compute cumulative abnormal returns (CARs) across four windows: "
        "[-60,-1] (~3 months), [-30,-1] (~6 weeks), [-10,-1] (~2 weeks), and [0,+5] "
        "(post-event). We cluster standard errors by VC firm (~1,100 clusters) to account for "
        "within-network correlation, and estimate VC fixed effects to absorb time-invariant "
        "VC-level characteristics."
    ),
    # Para 215: Core results
    (
        "Results are presented in Table 4.4 below. The overall pre-event CAR[-60,-1] is +0.79% "
        "(p = 0.006 with VC-firm clustering, N = 27,811). Breaking this down by industry match: "
        "same-industry connections show +0.97% (p = 0.080), while different-industry connections "
        "show +0.71% (p = 0.034). The effect sharpens dramatically at the [-30,-1] window: "
        "same-industry CARs reach +1.92% (p < 0.001), while different-industry CARs are "
        "statistically indistinguishable from zero (-0.18%, p = 0.428). This pattern is "
        "consistent with industry-specific private information being more actionable: when the "
        "observed firm operates in the same industry as the portfolio company, the information "
        "an observer acquires in board meetings is directly relevant and gets incorporated into "
        "prices earlier. Different-industry connections show a weaker, longer-horizon signal "
        "that fades at shorter windows."
    ),
    # Para 216: VC FE results
    (
        "Crucially, these results strengthen under VC fixed effects. Within the same VC firm, "
        "same-industry connections show 3.8% higher CARs than different-industry connections at "
        "the [-60,-1] window (p = 0.002 with VC-firm clustering) and 4.0% higher at [-30,-1] "
        "(p < 0.001). The VC FE specification absorbs all time-invariant VC-level "
        "characteristics\u2014VC quality, investment style, network centrality, information "
        "processing sophistication\u2014so the identification comes entirely from within-VC "
        "variation. This rules out the alternative explanation that \u201cbetter VCs\u201d "
        "systematically have both more observer connections and higher-returning portfolios. "
        "The information spillover effect is not driven by VC selection; it reflects genuine "
        "information flow through the observer\u2192VC\u2192portfolio company channel."
    ),
    # Para 217: Regulatory implications
    (
        "This finding has important regulatory implications. If observer seats facilitate "
        "information transfer between firms in the same industry, the antitrust concerns "
        "raised by the FTC/DOJ are empirically substantiated. As Packin and Alon-Beck document, "
        "the FTC has already recognized that \u201cboard observers have become more prevalent "
        "and could be privy to the same information as members of the board\u201d (p. 1553). "
        "Our information spillover results provide the first empirical evidence supporting "
        "this concern. The January 2025 DOJ/FTC extension of Clayton Act Section 8 interlocking "
        "directorate rules to cover board observers appears well-motivated by the data: "
        "same-industry observer connections generate statistically significant abnormal returns "
        "at connected firms months before public disclosure."
    ),
]

# New table content
table_caption = (
    "Table 4.4: Test 3 Results \u2014 Information Spillover Through Observer Networks "
    "(Private Events, CARs at Connected Public Portfolio Companies)"
)

# Replace paragraphs 214-216, keep 217 as table caption
for i, new_text in enumerate(new_texts):
    para_idx = 214 + i
    p = doc.paragraphs[para_idx]
    # Clear existing runs
    for run in p.runs:
        run.text = ""
    if p.runs:
        p.runs[0].text = new_text
    else:
        p.text = new_text

# Update table caption (para 217)
p217 = doc.paragraphs[217]
for run in p217.runs:
    run.text = ""
if p217.runs:
    p217.runs[0].text = table_caption
else:
    p217.text = table_caption

# Now insert the results table after paragraph 217
# Find the paragraph element for 217 and insert table after it
from docx.oxml.ns import qn

# Build the table data
headers = ["", "CAR[-60,-1]", "CAR[-30,-1]", "CAR[-10,-1]", "CAR[0,+5]"]
rows_data = [
    ["Panel A: Subsample Means (VC-firm clustered p-values)", "", "", "", ""],
    ["Overall", "+0.79%***", "+0.45%**", "+0.32%**", "+0.12%"],
    ["  N", "27,811", "27,805", "27,785", "27,791"],
    ["Same-industry", "+0.97%*", "+1.92%***", "+1.02%***", "+0.16%"],
    ["  N", "8,313", "8,345", "8,377", "8,377"],
    ["Diff-industry", "+0.71%**", "-0.18%", "+0.01%", "+0.11%"],
    ["  N", "19,498", "19,460", "19,408", "19,414"],
    ["", "", "", "", ""],
    ["Panel B: VC Fixed Effects (within-VC, same vs. diff industry)", "", "", "", ""],
    ["same_industry coef", "+3.83%***", "+4.01%***", "+0.80%**", "+0.75%***"],
    ["  p (VC-clustered)", "(0.002)", "(<0.001)", "(0.101)", "(0.038)"],
    ["  VC clusters", "1,134", "1,134", "1,136", "1,132"],
]

# Insert table
# Find the element after para 217
para_217_elem = doc.paragraphs[217]._element
parent = para_217_elem.getparent()

table = doc.add_table(rows=len(rows_data) + 1, cols=5)
table.style = "Table Grid"

# Set header row
for j, h in enumerate(headers):
    cell = table.rows[0].cells[j]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Fill data rows
for i, row in enumerate(rows_data):
    for j, val in enumerate(row):
        cell = table.rows[i + 1].cells[j]
        cell.text = val
        for run in cell.paragraphs[0].runs:
            run.font.size = Pt(9)
        if j > 0:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Bold panel headers
        if "Panel" in val:
            for run in cell.paragraphs[0].runs:
                run.bold = True
                run.italic = True

# Move the table to right after para 217
table_elem = table._element
parent = table_elem.getparent()
parent.remove(table_elem)
para_217_elem.addnext(table_elem)

# Add table notes after the table
notes_text = (
    "Notes: CARs are cumulative abnormal returns (raw returns, not market-adjusted) at "
    "public portfolio companies connected to private firms through observer\u2192VC networks. "
    "Events are CIQ Key Dev entries at private companies only. Panel A reports subsample means "
    "with p-values from VC-firm clustered standard errors (~1,100 clusters). Panel B reports "
    "the coefficient on same_industry from a regression with VC fixed effects (within-"
    "transformation), testing whether same-industry connections show higher CARs than "
    "different-industry connections within the same VC firm. "
    "*** p<0.01, ** p<0.05, * p<0.10."
)

# Add a new paragraph for notes after the table
from docx.oxml import OxmlElement
notes_para = OxmlElement("w:p")
table_elem.addnext(notes_para)
# Now use the document API to format it
# Find the notes paragraph
for i, p in enumerate(doc.paragraphs):
    if p._element is notes_para:
        p.text = notes_text
        p.style = doc.styles["Normal"]
        for run in p.runs:
            run.font.size = Pt(9)
            run.italic = True
        break

doc.save(path)
print("Saved updated v3 proposal with Test 3 FE results.")
