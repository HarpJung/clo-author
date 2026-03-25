"""Edit v3: Add structured NVCA survey statistics table in Section 2A."""

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = docx.Document("C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx")

# Find the paragraph right after "Furthermore, all entities surveyed which identified as traditional VCs..."
# That's the last paragraph of the NVCA Survey subsection, before "Legal Framework"
target_idx = None
for i, para in enumerate(doc.paragraphs):
    if "all entities surveyed which identified as traditional VCs" in para.text:
        target_idx = i
        print(f"Found traditional VCs paragraph at {i}")
        break

if target_idx is None:
    print("ERROR: Could not find target paragraph")
    exit()

body = doc.element.body
insert_after_elem = doc.paragraphs[target_idx]._element

# Add a summary paragraph
summary = doc.add_paragraph()
summary.style = doc.styles["Normal"]
run = summary.add_run(
    "The following table consolidates the key statistics from the NVCA Q4 2023 CFO "
    "Working-Group Survey as reported by Packin and Alon-Beck (2025). This survey, "
    "conducted in January 2024 among NVCA member firms, represents the first systematic "
    "empirical data on board observer prevalence in the VC/startup ecosystem. The survey "
    "was sent to approximately 200 member firms with a response rate of 42 to 66 firms; "
    "46 entities responded to the board observer questions (p. 1531)."
)
run.font.name = "Times New Roman"
run.font.size = Pt(12)
summary.paragraph_format.line_spacing = 2.0
summary.paragraph_format.space_after = Pt(6)
summary_elem = summary._element
body.remove(summary_elem)
insert_after_elem.addnext(summary_elem)

# Create the statistics table
table = doc.add_table(rows=11, cols=3)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Statistic", "Finding", "Source (Packin & Alon-Beck, 2025)"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

data = [
    ["Overall adoption rate",
     "82% of surveyed VC/AVC entities use board observers (38 of 46 respondents)",
     "p. 1514, 1531"],
    ["Non-adopters",
     "8 of 46 entities do not use board observers",
     "p. 1531"],
    ["Plans to increase",
     "8 entities plan to increase number of observers; 33 plan to keep current number",
     "p. 1532"],
    ["AUM > $500M adoption",
     "100% of entities with AUM exceeding $500 million use board observers",
     "p. 1515, 1534"],
    ["AUM breakdown (>$500M)",
     "$500M\u2013$1B: 10 entities; $1B\u2013$3B: 11 entities; $3B\u2013$5B: 5 entities; >$5B: 9 entities",
     "p. 1534"],
    ["AUM < $500M adoption",
     "Only 3 of 11 entities with AUM under $500M use board observers; 8 do not",
     "p. 1534"],
    ["Portfolio coverage (high)",
     "10 entities report >75% of portfolio companies have board observers",
     "p. 1533"],
    ["Portfolio coverage (moderate)",
     "16 entities report 25\u201375% of portfolio companies have observers",
     "p. 1533"],
    ["Portfolio coverage (low)",
     "5 entities report <25% of portfolio companies have observers; 7 did not disclose",
     "p. 1533"],
    ["Traditional VC adoption",
     "All entities identifying as traditional VCs unanimously use board observers (also identified as growth equity)",
     "p. 1535"],
]

for r_idx, row_data in enumerate(data):
    for c_idx, cell_text in enumerate(row_data):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"

# Move table after summary
table_elem = table._element
body.remove(table_elem)
summary_elem.addnext(table_elem)

# Add source note after table
note = doc.add_paragraph()
note.style = doc.styles["Normal"]
run = note.add_run(
    "Source: NVCA Q4 2023 CFO Working-Group Survey, as reported in Packin & Alon-Beck, "
    "\"Board Observers,\" 2025 U. Ill. L. Rev. 1507 (2025), pp. 1514\u20131535. The raw survey "
    "instrument and data are not publicly available; the NVCA provided access to the authors "
    "as acknowledged in the paper (p. 1531, fn. 165\u2013166). Survey conducted January 2024 among "
    "NVCA CFO Working-Group member firms."
)
run.font.name = "Times New Roman"
run.font.size = Pt(10)
run.italic = True
note.paragraph_format.space_after = Pt(12)
note_elem = note._element
body.remove(note_elem)
table_elem.addnext(note_elem)

# Save
output = "C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx"
doc.save(output)
print(f"Saved updated v3 to: {output}")
