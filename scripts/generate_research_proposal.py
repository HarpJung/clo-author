"""Generate Word document: Research Proposal — Three-Tier Board Governance."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 2.0  # double-spaced


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def bold_para(bold_text, normal_text=""):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
                run.font.name = 'Times New Roman'
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Times New Roman'
    return table


# =====================================================================
# TITLE PAGE
# =====================================================================
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run(
    "The Three-Tier Board:\n"
    "Directors, Observers, and the Accountability Gap\n"
    "in Venture-Backed Governance"
)
run.bold = True
run.font.size = Pt(16)
run.font.name = 'Times New Roman'

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Research Proposal")
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run("Harp Jung\nHarvard University\n\nMarch 2026")
run.font.size = Pt(12)

doc.add_page_break()

# =====================================================================
# 1. INTRODUCTION & MOTIVATION
# =====================================================================
heading("1. Introduction", 1)

doc.add_paragraph(
    "Board observers are present in 82% of venture capital-backed companies (NVCA 2024), "
    "yet no accounting, finance, or strategy journal has studied their governance role. "
    "This paper proposes the first systematic empirical examination of the three-tier board "
    "governance architecture\u2014directors, observers, and advisors\u2014that characterizes "
    "most institutionally-backed private firms."
)

doc.add_paragraph(
    "The three tiers differ fundamentally in accountability. Directors bear fiduciary duties "
    "and face personal liability; observers attend all board meetings and receive all materials "
    "but owe no fiduciary duties and bear no liability; advisors influence governance through "
    "social capital and expertise without any formal role. This structure creates a governance "
    "configuration where the participants with the most information (observers) may have the "
    "least responsibility, while the participants with legal duties (directors) may face conflicts "
    "that prevent effective monitoring."
)

doc.add_paragraph(
    "I exploit two quasi-natural experiments to test whether the accountability structure of "
    "board observers affects governance outcomes. First, in 2020, the National Venture Capital "
    "Association (NVCA) removed language from its model Investor Rights Agreement requiring "
    "observers to act \u201cin a fiduciary manner,\u201d eliminating the last contractual accountability "
    "anchor for observers. Second, in January 2025, the DOJ and FTC argued that board observers "
    "are subject to Clayton Act Section 8 interlocking directorate rules, creating a sudden "
    "increase in the regulatory cost of maintaining observer seats at competing firms."
)

doc.add_paragraph(
    "Together, these settings answer two complementary questions: Does removing the accountability "
    "anchor for observers worsen governance outcomes (NVCA 2020)? And does removing observers "
    "entirely reveal their governance value (Clayton Act 2025)?"
)

doc.add_page_break()

# =====================================================================
# 2. LITERATURE & CONTRIBUTION
# =====================================================================
heading("2. Related Literature and Contribution", 1)

heading("2.1 Board of Directors", 2)
doc.add_paragraph(
    "The governance role of boards of directors is extensively studied. Hermalin and Weisbach "
    "(2003) and Adams, Hermalin, and Weisbach (2010) survey the literature on board composition, "
    "independence, and monitoring effectiveness. Adams and Ferreira (2007) model the tension between "
    "monitoring and advising functions, showing that \u201cfriendly\u201d boards that prioritize advising "
    "may be optimal when the cost of reducing information flow outweighs the monitoring benefit. "
    "Fama and Jensen (1983) provide the foundational agency framework for understanding how boards "
    "reduce agency costs arising from the separation of ownership and control."
)

heading("2.2 Venture Capital Governance", 2)
doc.add_paragraph(
    "Kaplan and Stromberg (2003) document the complex system of cash flow, board, voting, and "
    "liquidation rights in VC contracts. Ewens and Malenko (2025) track board dynamics over "
    "the startup lifecycle using Form D filings, showing that boards evolve from entrepreneur-controlled "
    "at formation to VC-controlled at later stages, with independent directors serving mediation "
    "and then advising roles. Lerner (1995) establishes that VC board representation increases "
    "around key events. Broughman (2010, 2013) studies independent directors in VC-backed firms."
)

heading("2.3 Board Observers", 2)
doc.add_paragraph(
    "Packin and Alon-Beck (2025) provide the first academic treatment of board observers, "
    "documenting their legal status, prevalence (82% of VC entities), and the regulatory landscape "
    "including antitrust (Clayton Act Section 8), CFIUS, and ERISA implications. Their analysis "
    "is legal and institutional; they explicitly call for empirical research on governance consequences. "
    "Amornsiripanitch, Gompers, and Xuan (2019) report observer industry experience data as a "
    "secondary finding (38% for observers vs. 65% for directors) but do not study observer governance. "
    "No accounting, finance, or strategy paper studies observers empirically."
)

heading("2.4 NVCA Standardization", 2)
doc.add_paragraph(
    "Bartlett, Gulati, and Liskow (2023) document the standardization of VC contracts through "
    "NVCA model document adoption, which grew from 3% of startup charters in 2004 to 85% by 2022. "
    "Adoption was driven by concentration of legal services among approximately six specialist law "
    "firms. Their cosine similarity methodology for measuring template adoption provides a foundation "
    "for identifying which firms adopted the 2020 NVCA revisions. Critically, their study examines "
    "charters (publicly filed with Delaware); the Investor Rights Agreement containing observer "
    "provisions is a separate, generally private document that becomes public only when filed as an "
    "S-1 exhibit at IPO."
)

heading("2.5 D&O Insurance and Director Liability", 2)
doc.add_paragraph(
    "Core (2000), Lin, Officer, and Zou (2011), and Lin, Officer, Schmid, and Zou (2019) study "
    "how D&O insurance affects governance through moral hazard\u2014insulated directors may monitor "
    "less. Bradley and Chen (2011) connect director liability protection to the cost of debt. "
    "This literature focuses exclusively on directors; the interaction between D&O coverage and "
    "the observer/advisor tiers is unstudied."
)

heading("2.6 Interlocking Directorates", 2)
doc.add_paragraph(
    "Donelson, Hutzler, and Rhodes (2025) study the October 2022 DOJ enforcement wave against "
    "interlocking directors at PE/VC portfolio companies, finding that forced director departures "
    "weakened governance (less industry expertise, lower monitoring quality). Their event study "
    "methodology and sample construction provide a direct template for extending the analysis to "
    "observer interlocks following the January 2025 DOJ/FTC statement. Poberejsky (2025) finds "
    "that horizontal director interlocks reduce competition but increase innovation."
)

heading("2.7 Contribution", 2)
doc.add_paragraph("This paper contributes to the literature in five ways:")
bullet("First systematic empirical study of the three-tier governance architecture "
       "(directors + observers + advisors) as an integrated system.", "1. ")
bullet("First test of whether observer accountability structures affect governance outcomes, "
       "using the 2020 NVCA fiduciary language removal as a quasi-natural experiment.", "2. ")
bullet("First extension of the interlocking directorate literature to observer interlocks, "
       "using the January 2025 DOJ/FTC Clayton Act statement.", "3. ")
bullet("First evidence on whether observers create information bridges across VC portfolios, "
       "testing the mechanism underlying the DOJ/FTC\u2019s antitrust concern about observer-based "
       "competitive interlocks.", "4. ")
bullet("Construction of the first comprehensive dataset of board observer arrangements, "
       "combining S-1 filings, Capital IQ, PitchBook, and AlphaSense.", "5. ")

doc.add_page_break()

# =====================================================================
# 3. INSTITUTIONAL BACKGROUND
# =====================================================================
heading("3. Institutional Background", 1)

heading("3.1 The Three-Tier Governance Architecture", 2)

doc.add_paragraph(
    "When a venture capital or private equity firm invests in a company, the governance "
    "architecture that emerges typically includes three tiers of board-level participants:"
)

add_table(
    ["Tier", "Role", "Fiduciary Duty", "Voting Rights", "Information Access", "Personal Liability"],
    [
        ["1. Directors", "Formal board member", "Yes", "Yes", "Full", "Yes (mitigated by D&O)"],
        ["2. Observers", "Nonvoting board attendee", "No", "No", "Full (same as directors)", "No"],
        ["3. Advisors", "Informal advisor to board", "No", "No", "Varies (often limited)", "No"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "Observer rights are typically established in the Investor Rights Agreement (IRA) negotiated "
    "at each funding round. The IRA grants specific investors the right to designate a representative "
    "to attend all board meetings in a nonvoting observer capacity and to receive all notices, "
    "minutes, consents, and other materials provided to directors. Observer rights are often "
    "contingent on ownership thresholds (ranging from 4% to 25% of outstanding shares in our S-1 "
    "sample) and typically terminate upon the company\u2019s IPO, largely due to Regulation Fair "
    "Disclosure (Reg FD) concerns\u2014observers with access to material nonpublic information at a "
    "public company create selective disclosure risk."
)

heading("3.2 The 2020 NVCA Model Document Change", 2)

doc.add_paragraph(
    "The National Venture Capital Association (NVCA) publishes model legal documents\u2014standardized "
    "templates used in approximately 85% of VC financing transactions by 2022 (Bartlett, Gulati, "
    "and Liskow 2023). The NVCA is a trade association with no regulatory authority; adoption is "
    "voluntary but near-universal due to cost reduction, coordination benefits, and concentration "
    "of legal services among specialist law firms."
)

doc.add_paragraph(
    "Prior to 2020, the NVCA model Investor Rights Agreement included the following language "
    "regarding board observers:"
)

quote = doc.add_paragraph()
quote.style = 'Intense Quote'
quote.add_run(
    "The observer shall agree to hold in confidence and trust and to act in a fiduciary manner "
    "with respect to all information so provided."
)

doc.add_paragraph(
    "This language was legally anomalous\u2014observers have no fiduciary duties under Delaware law\u2014"
    "but it created a contractual behavioral anchor. In 2020, the NVCA deleted this language as a "
    "\u201ccorrection\u201d to align the template with legal reality. Post-2020 IRAs rely solely on "
    "separate confidentiality provisions (NDAs, side letters) for observer obligations. The removal "
    "was driven by legal accuracy concerns across the industry, not by governance failures at "
    "specific firms, supporting the exogeneity of the treatment."
)

doc.add_paragraph(
    "Direct evidence of this language variation appears in S-1 filings. For example, Coinbase\u2019s "
    "pre-2020 IRA (Exhibit 4.2) requires observers to act \u201cin a fiduciary manner,\u201d while "
    "Sera Prognostics\u2019 2021 IRA explicitly states that the observer \u201cshall not\u2026have or be "
    "deemed to have any fiduciary or other duties.\u201d"
)

heading("3.3 The January 2025 DOJ/FTC Clayton Act Statement", 2)

doc.add_paragraph(
    "On January 10, 2025, the DOJ and FTC filed a joint Statement of Interest in Musk v. Altman "
    "(N.D. Cal.) arguing that board observers are subject to Section 8 of the Clayton Act\u2019s "
    "prohibition on interlocking directorates. Prior to this statement, the FTC had explicitly "
    "excluded observers from Section 8 coverage. The statement was adopted unanimously, including "
    "by the incoming administration\u2019s FTC designee, suggesting bipartisan durability."
)

doc.add_paragraph(
    "This dramatically increased the regulatory cost of observer arrangements at competing firms. "
    "VC/PE firms with portfolio companies in overlapping markets faced sudden antitrust exposure "
    "for observer seats that had previously been outside Section 8\u2019s reach. Donelson, Hutzler, "
    "and Rhodes (2025) document that similar enforcement against director interlocks in October 2022 "
    "led to director departures, loss of industry expertise, and weakened governance."
)

doc.add_page_break()

# =====================================================================
# 4. HYPOTHESES
# =====================================================================
heading("4. Hypotheses", 1)

heading("4.1 NVCA 2020: Observer Accountability and Governance Outcomes", 2)

doc.add_paragraph(
    "The removal of fiduciary language from the NVCA model IRA reduces the contractual "
    "accountability anchor for observers. If this anchor influenced observer behavior at the "
    "margin\u2014even symbolically\u2014its removal should be associated with weaker governance."
)

bold_para("H1 (Financial Reporting Quality): ",
    "Firms whose IRAs lack fiduciary language for observers (post-2020 NVCA) exhibit lower "
    "financial reporting quality post-IPO, as measured by higher audit fees, more restatements, "
    "and greater discretionary accruals, relative to firms with fiduciary language (pre-2020).")

bold_para("H2 (Monitoring Intensity): ",
    "Firms with post-2020 IRAs exhibit weaker CEO turnover-performance sensitivity post-IPO, "
    "consistent with observers being less engaged as governance monitors when freed from the "
    "fiduciary anchor.")

bold_para("H3 (Investor Pricing): ",
    "Follow-on investors in private firms funded under post-2020 NVCA terms demand higher "
    "returns (lower valuation step-ups) or stronger protective provisions, reflecting the "
    "market\u2019s assessment of weaker observer accountability.")

bold_para("H4 (Observer-to-Director Transitions): ",
    "Observers under post-2020 IRAs are less likely to transition to full director roles, "
    "consistent with the fiduciary language serving as a stepping stone toward greater "
    "governance commitment.")

heading("4.2 Clayton Act 2025: Observer Value Revealed by Removal", 2)

doc.add_paragraph(
    "The DOJ/FTC statement increases the cost of maintaining observers at competing firms. "
    "If observers provided governance value (monitoring, information, expertise), their "
    "forced removal should worsen outcomes."
)

bold_para("H5 (Observer Departures): ",
    "After January 2025, observer seats at firms with competitive overlap decline as "
    "observers resign or firms restructure to eliminate interlocks.")

bold_para("H6 (Governance Consequences of Removal): ",
    "Firms that lose observers due to the regulatory shock experience worse governance "
    "outcomes\u2014more down rounds, longer time to next funding, higher failure rates\u2014"
    "relative to firms that retain observers.")

bold_para("H7 (Substitution to Advisors): ",
    "Investors who lose observer seats at competitors substitute into advisory board roles, "
    "which are further from Section 8\u2019s reach. The three-tier system adapts: regulatory "
    "pressure on Tier 2 (observers) pushes activity to Tier 3 (advisors).")

heading("4.3 Information Leakage: Observer-Level Tests", 2)

doc.add_paragraph(
    "The NVCA fiduciary language specifically addressed confidential information handling. "
    "If removing the fiduciary anchor makes observers less careful with information, this should "
    "be detectable at the observer level\u2014not just at the firm level. Observers create potential "
    "information bridges between portfolio companies of the same VC/PE investor: an observer at "
    "Company A who reports to VC firm X may (consciously or not) transmit information that "
    "benefits VC X\u2019s other portfolio companies."
)

bold_para("H8 (Cross-Portfolio Return Spillovers): ",
    "Material events at Company A (where VC X has an observer) generate abnormal returns at "
    "VC X\u2019s other public portfolio companies in the same industry, consistent with information "
    "flowing through the observer channel. Spillovers through observer connections are larger "
    "than spillovers through director connections, reflecting the accountability gap: observers "
    "face no fiduciary constraint on information sharing while directors do.")

bold_para("H9 (NVCA Interaction): ",
    "Cross-portfolio return spillovers through observer connections increase after the 2020 "
    "NVCA fiduciary language removal, consistent with the behavioral anchor having constrained "
    "information flow prior to its removal.")

bold_para("H10 (Pre-Announcement Information Leakage): ",
    "At public firms with persistent board observers (post-IPO), abnormal trading volume and "
    "price drift are larger in the days preceding material announcements (M&A, earnings) "
    "relative to matched firms without observers, consistent with observers creating informed "
    "trading opportunities.")

doc.add_page_break()

# =====================================================================
# 5. DATA
# =====================================================================
heading("5. Data and Sample Construction", 1)

heading("5.1 Data Sources", 2)

doc.add_paragraph(
    "I construct the dataset using four complementary sources, each providing a different "
    "layer of information about board governance architecture:"
)

add_table(
    ["Layer", "Source", "What It Provides", "Coverage", "Access"],
    [
        ["1. Broad cross-section",
         "S&P Capital IQ (WRDS)",
         "Board observer/advisor presence; named individuals with titles, current/former flags; observer-to-VC affiliation links for network construction",
         "2,602 private firms with observer data; 345,099 private firms with any board data; 461,599 total firms",
         "SQL via WRDS"],
        ["2. Deep governance detail",
         "S-1 filings (SEC EDGAR)",
         "Full IRA text with observer provisions: ownership thresholds, fiduciary language, termination conditions, named observers",
         "1,133 S-1 files mentioning \"board observer\" (2020-2024); ~77+ unique companies in first 100 results",
         "EDGAR EFTS API + exhibit fetch"],
        ["3. Historical enrichment",
         "PitchBook",
         "Dedicated \"Board Members & Observers\" section with historical tracking; round-level valuations; exit data",
         "Broad VC/PE coverage; available on WRDS (pending access activation) or via web interface",
         "WRDS SQL (pending) or manual"],
        ["4. Document search",
         "AlphaSense",
         "AI-powered semantic search across 500M+ financial documents: SEC filings (S-1, DEF 14A, 10-K, 8-K), earnings calls, broker research, expert transcripts (Tegus)",
         "All SEC filings; earnings call transcripts; broker research",
         "Web (Harvard subscription)"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "AlphaSense supplements the structured databases by enabling rapid identification of "
    "observer-related language across all SEC filing types\u2014including proxy statements (DEF 14A) "
    "and annual reports (10-K) that may reference observer arrangements for public firms beyond "
    "the S-1 sample. It also provides access to expert transcripts (via Tegus) containing "
    "practitioner discussions of observer governance."
)

doc.add_paragraph()

heading("5.2 Treatment Variable: NVCA Fiduciary Language", 2)

doc.add_paragraph(
    "For the NVCA test, the treatment variable is the presence or absence of fiduciary language "
    "in the Investor Rights Agreement governing observer obligations."
)

bold_para("S-1 subsample (direct coding): ",
    "I read the IRA exhibit filed with each S-1 and code whether the agreement contains "
    "pre-2020 fiduciary language (\u201cact in a fiduciary manner\u201d) or post-2020 language "
    "(no fiduciary reference, or explicit \u201cno fiduciary duty\u201d clause). A feasibility study "
    "of 18 VC/PE-backed IPOs confirmed this language is clearly identifiable and varies "
    "across firms: Coinbase\u2019s pre-2020 IRA contains fiduciary language; Sera Prognostics\u2019 "
    "2021 IRA explicitly disclaims fiduciary duty.")

bold_para("Full sample (proxy): ",
    "For the broader CIQ/PitchBook sample of private firms, I proxy for NVCA IRA adoption "
    "using the firm\u2019s funding date relative to the 2020 template change. Firms whose most "
    "recent funding round closed after 2020 are classified as likely post-2020 IRA adopters, "
    "given the 85% NVCA adoption rate (Bartlett et al. 2023). Staggered adoption within years "
    "provides additional variation.")

heading("5.3 Treatment Variable: Clayton Act Competitive Interlocks", 2)

doc.add_paragraph(
    "For the Clayton Act test, the treatment is the January 10, 2025 DOJ/FTC statement. "
    "The treatment group consists of firms where the observer\u2019s sponsoring investor also holds "
    "stakes in competing companies (observer-based competitive interlocks). The control group "
    "consists of firms with observers whose investors do not have competing portfolio companies. "
    "Competitive overlap is identified using industry classification (SIC/NAICS) and, where "
    "available, the text-based network industry classification (TNIC) of Hoberg and Phillips, "
    "following the methodology of Donelson, Hutzler, and Rhodes (2025)."
)

heading("5.4 Outcome Variables", 2)

doc.add_paragraph(
    "I employ outcome variables across three categories, prioritized by data feasibility. "
    "Categories 1 and 2 measure firm-level governance outcomes. Category 3 measures "
    "observer-level information leakage\u2014the most direct test of the confidential information "
    "mechanism underlying the NVCA fiduciary language."
)

bold_para("Category 1: Post-IPO firm-level outcomes (highest feasibility\u2014all available via WRDS SQL):")
add_table(
    ["Variable", "Database", "Measure"],
    [
        ["Audit fees", "Audit Analytics", "Log(audit fees), years 1-3 post-IPO"],
        ["Restatements", "Audit Analytics", "Binary indicator within 3 years of IPO"],
        ["Internal control weaknesses", "Audit Analytics", "SOX 302/404 material weakness"],
        ["CEO/CFO turnover", "Audit Analytics", "Binary indicator of executive change"],
        ["IPO underpricing", "CRSP", "First-day return"],
        ["Post-IPO return volatility", "CRSP", "Standard deviation of daily returns, year 1"],
    ]
)

doc.add_paragraph()
bold_para("Category 2: Private firm-level outcomes (high feasibility\u2014CIQ via SQL, supplemented by PitchBook):")
add_table(
    ["Variable", "Database", "Measure"],
    [
        ["Company failure", "CIQ", "Status = Out of Business or Liquidating"],
        ["CEO/founder replacement", "CIQ Key Dev + Professional", "Executive change event or title transition"],
        ["Observer-to-director transitions", "CIQ Professional", "Title change from Board Observer to Director"],
        ["Exit quality", "CIQ + PitchBook", "IPO vs. acquisition vs. failure"],
        ["Down rounds", "PitchBook", "Valuation decrease at next funding round"],
        ["Time to next round", "PitchBook", "Days between consecutive funding events"],
    ]
)

doc.add_paragraph()
bold_para("Category 3: Observer-level information leakage outcomes (high feasibility\u2014CIQ network + CRSP/TAQ via WRDS):")
doc.add_paragraph(
    "These tests exploit the observer network constructible from CIQ: each observer\u2019s "
    "professional records link them to their VC/PE firm affiliation, and from there to the "
    "VC\u2019s other portfolio companies. For public portfolio companies in this network, market "
    "data from CRSP and TAQ enables direct measurement of information flows."
)
add_table(
    ["Variable", "Database", "Measure"],
    [
        ["Cross-portfolio return spillovers",
         "CIQ (network) + CRSP",
         "CARs at VC X\u2019s public portfolio companies around material events at the observed firm; "
         "compare observer-connected vs. director-connected spillovers"],
        ["Pre-announcement price drift",
         "CRSP",
         "Abnormal returns in [-10, -1] window before M&A/earnings announcements at public firms with observers vs. matched firms without"],
        ["Pre-announcement abnormal volume",
         "CRSP / TAQ",
         "Abnormal trading volume in days preceding material events at firms with persistent post-IPO observers"],
        ["Bid-ask spread / PIN",
         "TAQ",
         "Information asymmetry measures at public firms with observers; higher spreads indicate more informed trading"],
        ["Observer departure information effects",
         "CIQ + CRSP",
         "Change in information asymmetry measures when an observer departs (Clayton Act forced removal)"],
    ]
)

doc.add_paragraph()

heading("5.5 Sample Construction", 2)

bold_para("NVCA test\u2014S-1 subsample: ",
    "All VC-backed IPOs on U.S. exchanges, 2017-2024. The pre-period (2017-2019) captures firms "
    "funded under pre-2020 NVCA terms; the post-period (2021-2024) captures firms increasingly "
    "funded under post-2020 terms. IRAs are coded directly from S-1 exhibits. EDGAR full-text "
    "search identifies 1,133 S-1 files mentioning \u201cboard observer\u201d in 2020-2024 alone, "
    "suggesting a sample of several hundred unique firms.")

bold_para("NVCA test\u2014full sample: ",
    "All VC/PE-backed private firms in CIQ with board observer data (2,602 firms), supplemented "
    "by PitchBook. Treatment is proxied by most recent funding date relative to 2020.")

bold_para("Clayton Act test: ",
    "All firms in CIQ/PitchBook with observer-based competitive interlocks as of December 2024. "
    "Pre-period: 2023-2024. Post-period: January 2025 onward. Event study design following "
    "Donelson et al. (2025).")

doc.add_page_break()

# =====================================================================
# 6. EMPIRICAL STRATEGY
# =====================================================================
heading("6. Empirical Strategy", 1)

heading("6.1 NVCA 2020: Difference-in-Differences", 2)

doc.add_paragraph(
    "The baseline specification for the NVCA test is a difference-in-differences design "
    "comparing post-IPO outcomes for firms with pre-2020 vs. post-2020 IRA language:"
)

eq = doc.add_paragraph()
eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq.add_run("Y").italic = True
eq_run = eq.add_run("\u1d62\u209c = \u03b1 + \u03b2 \u00d7 PostNVCA")
eq.add_run("\u1d62").italic = True
eq.add_run(" + \u03b3 \u00d7 X")
eq.add_run("\u1d62\u209c").italic = True
eq.add_run(" + \u03b4").italic = True
eq.add_run("\u209c + \u03b5")
eq.add_run("\u1d62\u209c").italic = True

doc.add_paragraph(
    "where Y is the outcome variable (audit fees, restatements, etc.), PostNVCA is an indicator "
    "equal to one if the firm\u2019s IRA uses post-2020 language (no fiduciary anchor), X is a vector "
    "of firm-level controls (size, profitability, leverage, board size, VC reputation, number of "
    "funding rounds), and \u03b4 captures year fixed effects. The coefficient \u03b2 captures the "
    "differential governance outcome associated with the removal of the fiduciary accountability anchor."
)

doc.add_paragraph(
    "For the private firm sample, I replace PostNVCA with an indicator for whether the firm\u2019s "
    "most recent funding round closed after 2020, exploiting staggered adoption timing for "
    "identification. Industry and vintage-year fixed effects absorb common shocks."
)

heading("6.2 Cross-Portfolio Return Spillovers (H8-H9)", 2)

doc.add_paragraph(
    "The information leakage test uses the observer network constructed from CIQ to measure "
    "whether material events at one portfolio company generate abnormal returns at another "
    "portfolio company connected through a common observer. The baseline specification is:"
)

eq2 = doc.add_paragraph()
eq2.alignment = WD_ALIGN_PARAGRAPH.CENTER
eq2.add_run(
    "CAR(Company B) = \u03b1 + \u03b2\u2081 \u00d7 Event(A) \u00d7 CommonObserver "
    "+ \u03b2\u2082 \u00d7 Event(A) \u00d7 CommonDirector + \u03b3X + \u03b5"
)

doc.add_paragraph(
    "where CAR(Company B) is the cumulative abnormal return at Company B (a public portfolio "
    "company of VC X) around a material event at Company A (where VC X has a board observer), "
    "CommonObserver is an indicator for the observer connection, CommonDirector is an indicator "
    "for a director (but not observer) connection at the same VC, and X is a vector of controls "
    "including industry relatedness, firm size, and VC characteristics."
)

doc.add_paragraph(
    "The coefficient \u03b2\u2081 captures the information spillover through the observer channel. "
    "If \u03b2\u2081 > \u03b2\u2082, observers create more information leakage than directors\u2014consistent "
    "with the accountability gap hypothesis, since directors face fiduciary constraints on "
    "information sharing while observers do not. To test H9 (NVCA interaction), I interact "
    "CommonObserver with a PostNVCA indicator, testing whether spillovers increase after the "
    "2020 fiduciary language removal."
)

doc.add_paragraph(
    "Network construction from CIQ is feasible: a preliminary analysis of 100 Board Observers "
    "found that their other professional positions include 315 records at Private Investment Firms "
    "(VC/PE affiliations) and 519 Director positions at other portfolio companies, confirming that "
    "the observer\u2192VC\u2192portfolio company links are identifiable in CIQ."
)

heading("6.3 Clayton Act 2025: Event Study", 2)

doc.add_paragraph(
    "For the Clayton Act test, I follow Donelson, Hutzler, and Rhodes (2025) in implementing "
    "an event study around the January 10, 2025 DOJ/FTC statement. For public firms with "
    "observer-based competitive interlocks, I estimate cumulative abnormal returns (CARs) in "
    "short windows around the event date. For private firms, I compare observer departure rates, "
    "board restructuring, and subsequent governance outcomes in treated firms (competitive interlocks) "
    "vs. control firms (observers without competitive overlap) in the months following the statement."
)

heading("6.4 Identification Concerns", 2)

bold_para("Endogeneity of observer arrangements: ",
    "Firms that have observers may differ from firms that do not in unobservable ways. The NVCA "
    "design mitigates this by comparing within the set of firms that have observers, exploiting "
    "variation in the fiduciary language (pre vs. post-2020) that is determined by template timing, "
    "not firm-level governance choices.")

bold_para("COVID confound: ",
    "The 2020 NVCA change coincides with COVID-19. Staggered adoption within the 2019-2022 "
    "window provides within-period variation. Additionally, I can control for COVID-era market "
    "conditions and industry-specific pandemic effects.")

bold_para("Selection into IPO: ",
    "The S-1 subsample is conditioned on IPO, a selected sample. The private firm sample "
    "(CIQ + PitchBook) does not condition on IPO and serves as a robustness check. If results "
    "hold in both samples, selection bias is unlikely to drive the findings.")

bold_para("Treatment strength: ",
    "A potential concern is that the fiduciary language was always legally meaningless, so its "
    "removal changes nothing. I address this by (a) testing for actual outcome differences, "
    "(b) examining whether investors price the language change (H3), and (c) citing practitioner "
    "evidence from interviews that contractual language shapes behavioral norms even when not "
    "legally binding.")

doc.add_page_break()

# =====================================================================
# 7. FEASIBILITY EVIDENCE
# =====================================================================
heading("7. Feasibility Evidence", 1)

doc.add_paragraph(
    "A preliminary feasibility study confirms that the proposed data strategy is viable:"
)

heading("7.1 S-1 Filing Analysis", 2)
doc.add_paragraph(
    "EDGAR full-text search identifies 1,133 files in S-1 filings (2020-2024) mentioning "
    "\u201cboard observer,\u201d with 77 unique companies in the first 100 results. A detailed "
    "analysis of 18 VC/PE-backed IPOs found observer provisions in 14 of 17 confirmed cases "
    "(82% prevalence). Observer arrangements are standardized enough for systematic coding but "
    "varied enough for cross-sectional analysis. Key coding variables include observer count, "
    "observer identity, ownership thresholds, fiduciary language (present/absent), and post-IPO "
    "termination conditions."
)

heading("7.2 Capital IQ (WRDS)", 2)
doc.add_paragraph(
    "S&P Capital IQ on WRDS contains 2,602 unique private companies with board observer records "
    "and 21,697 private companies with advisory board member records. Observer data is captured "
    "through the professional title field (e.g., \u201cBoard Observer,\u201d \u201cFormer Board Observer,\u201d "
    "\u201cNon-Voting Observer\u201d), with approximately 5,000 total records including current and former "
    "observers. The data is bulk-downloadable via SQL and linkable to CIQ\u2019s company universe "
    "(345,099 private firms with any board data)."
)

heading("7.3 PitchBook", 2)
doc.add_paragraph(
    "PitchBook maintains a dedicated \u201cBoard Members & Observers\u201d section with historical "
    "tracking of observer appointments and departures. Coverage includes round-level valuations, "
    "investor participation, and exit outcomes. PitchBook data is available on WRDS "
    "(pitchbk.personboardseatrelation table with roleonboard, startdate, enddate, and "
    "representingname fields), though access activation is pending. The web interface serves "
    "as a validation source with limited bulk extraction (10 rows/day for private data)."
)

heading("7.4 AlphaSense", 2)
doc.add_paragraph(
    "AlphaSense provides AI-powered semantic search across 500+ million financial documents, "
    "including all SEC filing types, earnings call transcripts, broker research, and Tegus "
    "expert interview transcripts. For this study, AlphaSense accelerates S-1 exhibit identification "
    "and enables expansion of the observer search to proxy statements (DEF 14A) and annual reports "
    "(10-K), potentially extending public firm observer coverage beyond the 431 companies identified "
    "in CIQ. AlphaSense is accessible via Harvard\u2019s institutional subscription."
)

heading("7.5 Observer Network Construction from CIQ", 2)
doc.add_paragraph(
    "A critical feasibility test confirmed that CIQ\u2019s professional records enable construction "
    "of the observer\u2192VC/PE firm\u2192portfolio company network required for the information leakage "
    "tests (H8-H10). For example, Martin Brand appears in CIQ as both a Board Observer at Bumble Inc. "
    "and Head of Blackstone Capital Partners at Blackstone Inc., with Director positions at 15+ other "
    "Blackstone portfolio companies. Across a sample of 100 observers, other positions include 315 "
    "records at Private Investment Firms (the VC/PE affiliation link) and 519 Director positions at "
    "portfolio companies (the cross-portfolio link). This confirms the observer network is identifiable "
    "and dense enough for the spillover analysis."
)

doc.add_page_break()

# =====================================================================
# 8. EXPECTED CONTRIBUTIONS
# =====================================================================
heading("8. Expected Contributions", 1)

bullet("Empirical: Construct the first comprehensive dataset of board observer arrangements "
       "across private and public firms, enabling future research on the two governance tiers "
       "(observers and advisors) currently absent from the literature.", "Data contribution: ")

bullet("Theoretical: Extend agency theory to a novel principal class\u2014observers as "
       "\u201chigh-information, low-duty\u201d governance participants\u2014and test whether "
       "accountability structures affect governance outcomes.", "Theory contribution: ")

bullet("Policy: Inform ongoing regulatory debates about observer status under antitrust law "
       "(Clayton Act Section 8), securities regulation (Reg FD), and corporate governance "
       "standards.", "Policy contribution: ")

bullet("Methodology: Demonstrate that S-1 filings, Capital IQ professional titles, and PitchBook "
       "board data can be combined to study private firm governance at scale\u2014a methodology "
       "applicable to other governance questions beyond observers.", "Methodological contribution: ")

doc.add_page_break()

# =====================================================================
# 9. TIMELINE
# =====================================================================
heading("9. Proposed Timeline", 1)

add_table(
    ["Phase", "Task", "Target"],
    [
        ["1", "Pull full CIQ observer and advisory board extracts from WRDS", "April 2026"],
        ["2", "Build S-1 IRA coding protocol; pilot on 18-company sample", "April 2026"],
        ["3", "Expand S-1 coding to full sample (~200-300 firms)", "May-June 2026"],
        ["4", "Merge with outcome databases (Audit Analytics, CRSP, Compustat)", "June 2026"],
        ["5", "PitchBook validation sample (50-100 firms)", "June 2026"],
        ["6", "NVCA DiD estimation", "July 2026"],
        ["7", "Clayton Act event study", "July-August 2026"],
        ["8", "Robustness tests and manuscript drafting", "August-September 2026"],
        ["9", "Conference presentation (target: AFA/EFA)", "October 2026"],
    ]
)

doc.add_page_break()

# =====================================================================
# REFERENCES
# =====================================================================
heading("References", 1)

refs = [
    "Adams, R. B., & Ferreira, D. (2007). A theory of friendly boards. Journal of Finance, 62(1), 217-250.",
    "Adams, R. B., Hermalin, B. E., & Weisbach, M. S. (2010). The role of boards of directors in corporate governance: A conceptual framework and survey. Journal of Economic Literature, 48(1), 58-107.",
    "Amornsiripanitch, N., Gompers, P. A., & Xuan, Y. (2019). More than money: Venture capitalists on boards. Journal of Law, Economics, and Organization, 35(3), 513-555.",
    "Bartlett, R. P., Gulati, M., & Liskow, E. (2023). Standardization and innovation in venture capital contracting. Journal of Law and Economics, 55(1).",
    "Bradley, M., & Chen, D. (2011). Corporate governance and the cost of debt. Journal of Corporate Finance, 17(1), 83-107.",
    "Broughman, B. J. (2010). The role of independent directors in startup firms. Utah Law Review, 2010(1).",
    "Core, J. E. (2000). The directors' and officers' insurance premium: An outside assessment of the quality of corporate governance. Journal of Law, Economics, and Organization, 16(2), 449-477.",
    "Donelson, D. C., Hutzler, A. D., & Rhodes, A. (2025). Does antitrust enforcement against interlocking directorates impair corporate governance? Journal of Accounting and Economics.",
    "Ewens, M., & Malenko, N. (2025). Board dynamics over the startup life cycle. Review of Financial Studies.",
    "Fama, E. F., & Jensen, M. C. (1983). Separation of ownership and control. Journal of Law and Economics, 26(2), 301-325.",
    "Hermalin, B. E., & Weisbach, M. S. (2003). Boards of directors as an endogenously determined institution: A survey of the economic literature. Economic Policy Review, 9(1), 7-26.",
    "Kaplan, S. N., & Stromberg, P. (2003). Financial contracting theory meets the real world: An empirical analysis of venture capital contracts. Review of Economic Studies, 70(2), 281-315.",
    "Lerner, J. (1995). Venture capitalists and the oversight of private firms. Journal of Finance, 50(1), 301-318.",
    "Lin, C., Officer, M. S., & Zou, H. (2011). Directors' and officers' liability insurance and acquisition outcomes. Journal of Financial Economics, 102(3), 507-525.",
    "Lin, C., Officer, M. S., Schmid, T., & Zou, H. (2019). Is skin in the game a game changer? Evidence from mandatory changes of D&O insurance policies. Journal of Accounting and Economics, 68(1).",
    "NVCA. (2024). Q4 2023 CFO Working-Group Survey. National Venture Capital Association.",
    "Packin, N. G., & Alon-Beck, A. (2025). Board observers. University of Illinois Law Review, 2025(5).",
    "Poberejsky, R. (2025). Interlocking directorates, competition, and innovation. Job market paper, Northwestern University.",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Inches(-0.5)
    p.paragraph_format.left_indent = Inches(0.5)
    for run in p.runs:
        run.font.size = Pt(11)

# =====================================================================
# SAVE
# =====================================================================
output_path = os.path.normpath(os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "..",
    "Research Proposal -- Three-Tier Board Governance.docx"
))
doc.save(output_path)
print(f"Saved to: {output_path}")
