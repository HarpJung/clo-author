"""Edit v3: Replace DoorDash and Rivian with OpenAI and Airbnb (have advisory boards)."""

import docx
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

doc = docx.Document("C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx")

# Find and replace the Rivian paragraph
for i, para in enumerate(doc.paragraphs):
    if para.text.startswith("Rivian (IPO 2021, Electric Vehicles)."):
        # Replace the entire paragraph
        for run in para.runs:
            run.text = ""
        # Add new content
        rb = para.add_run("OpenAI (2023\u20132024, Artificial Intelligence). ")
        rb.bold = True
        rb.font.name = "Times New Roman"
        rb.font.size = Pt(12)
        rn = para.add_run(
            "OpenAI provides the most publicly visible example of the three-tier architecture. "
            "Its board of directors includes CEO Sam Altman, along with directors such as "
            "retired U.S. Army General Paul Nakasone (former NSA Director), Nicole Seligman "
            "(former EVP and Global General Counsel of Sony), Fidji Simo (CEO and Chair of "
            "Instacart), and Sue Desmond-Hellmann (former CEO of the Bill and Melinda Gates "
            "Foundation). In November 2023, Microsoft \u2014 which had invested over $13 billion "
            "in OpenAI \u2014 was granted a nonvoting board observer seat, filled by Dee Templeton "
            "(a 25-year Microsoft veteran and head of Technology and Research Partnerships). "
            "Microsoft\u2019s observer arrangement drew intense regulatory scrutiny from the DOJ "
            "and FTC, ultimately leading Microsoft to voluntarily relinquish the seat in July "
            "2024. Separately, OpenAI established an advisory structure including former SEC "
            "Chair Jay Clayton and former CIA General Counsel Courtney Elwood as advisors. "
            "The OpenAI case illustrates all three tiers in action \u2014 and demonstrates the "
            "regulatory risks that Packin and Alon-Beck (2025) document."
        )
        rn.font.name = "Times New Roman"
        rn.font.size = Pt(12)
        print(f"Replaced Rivian with OpenAI at paragraph {i}")
        break

# Find and replace the DoorDash paragraph
for i, para in enumerate(doc.paragraphs):
    if para.text.startswith("DoorDash (IPO 2020, Food Delivery)."):
        for run in para.runs:
            run.text = ""
        rb = para.add_run("Airbnb (IPO 2020, Travel/Hospitality). ")
        rb.bold = True
        rb.font.name = "Times New Roman"
        rb.font.size = Pt(12)
        rn = para.add_run(
            "Airbnb\u2019s 9-member board at IPO included all three co-founders (Brian Chesky as "
            "CEO and Chairman, Nathan Blecharczyk as Chief Strategy Officer, and Joseph Gebbia "
            "as Chairman of Airbnb.org), VC directors Jeffrey Jordan (Managing Partner at "
            "Andreessen Horowitz, former CEO of OpenTable) and Alfred Lin (Partner at Sequoia "
            "Capital, former Chairman/COO of Zappos), and independent directors Angela Ahrendts "
            "(former SVP of Apple Retail and former CEO of Burberry), Kenneth Chenault (former "
            "Chairman/CEO of American Express), Ann Mather (former CFO of Pixar), and Belinda "
            "Johnson (former Airbnb COO). The S-1 references an Investors\u2019 Rights Agreement "
            "with observer provisions that terminated at IPO. Uniquely, Airbnb established both "
            "a Stakeholder Committee (an advisory committee of the board \u201cadvisory in nature,\u201d "
            "focused on balancing host, guest, community, and employee interests) and a Host "
            "Advisory Board (approximately 15 prominent Airbnb hosts providing operational "
            "feedback), seeded with a 9.2 million share Host Endowment Fund. This illustrates "
            "how the \u201cadvisor\u201d tier can take multiple forms: a governance-level advisory "
            "committee and a stakeholder-level operational advisory body \u2014 neither with voting "
            "power, fiduciary duties, or formal board authority."
        )
        rn.font.name = "Times New Roman"
        rn.font.size = Pt(12)
        print(f"Replaced DoorDash with Airbnb at paragraph {i}")
        break

# Update the summary table - find it and update the Rivian and DoorDash rows
for table in doc.tables:
    for row in table.rows:
        cell0_text = row.cells[0].text.strip()
        if "Rivian" in cell0_text:
            row.cells[0].text = "OpenAI\n(2023-24)"
            row.cells[1].text = "~7 directors: CEO Altman, Gen. Nakasone (fmr NSA Dir.), Seligman (fmr Sony GC), Simo (Instacart CEO), Desmond-Hellmann (fmr Gates Foundation CEO)"
            row.cells[2].text = "1 observer: Microsoft (Dee Templeton, 25-yr MSFT veteran). Relinquished Jul 2024 under DOJ/FTC scrutiny."
            row.cells[3].text = "Advisory structure incl. fmr SEC Chair Jay Clayton, fmr CIA GC Courtney Elwood"
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
                        run.font.name = "Times New Roman"
            print("Updated OpenAI row in table")

        if "DoorDash" in cell0_text:
            row.cells[0].text = "Airbnb\n(2020)"
            row.cells[1].text = "9 directors: 3 co-founders, 2 VC directors (a16z, Sequoia), 4 independent (fmr Apple SVP, fmr AmEx CEO, fmr Pixar CFO, fmr Airbnb COO)"
            row.cells[2].text = "Per IRA: VC co-investors as observers. All terminated at IPO."
            row.cells[3].text = "Stakeholder Committee (advisory, governance-level) + Host Advisory Board (15 hosts, operational) + Host Endowment (9.2M shares)"
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        run.font.size = Pt(8)
                        run.font.name = "Times New Roman"
            print("Updated Airbnb row in table")

# Update the interpretive paragraph about "only Coinbase" having advisory
for i, para in enumerate(doc.paragraphs):
    if "only Coinbase established one" in para.text:
        for run in para.runs:
            run.text = ""
        para.add_run(
            "Several patterns emerge from these examples. First, directors are consistently founders, "
            "lead VC partners, and experienced independent directors with deep operational expertise "
            "(former CEOs, CFOs, and general counsels) \u2014 people who have committed to fiduciary "
            "responsibility. Second, observers are typically additional VC investors (co-investors "
            "who did not receive a board seat), strategic investors (Microsoft at OpenAI, Amazon at "
            "other companies), or PE sponsors maintaining parallel information channels (Blackstone "
            "at Bumble). Third, formal advisory structures vary widely: Coinbase established a "
            "post-IPO Global Advisory Council with regulatory figures (former SEC Chair, former "
            "Senator, former Fed President); Airbnb created both a governance-level Stakeholder "
            "Committee and an operational Host Advisory Board; OpenAI has advisors including former "
            "SEC and CIA officials; while Reddit and Bumble disclosed no formal advisory boards in "
            "their S-1 filings. This confirms the pattern from our interviews: advisory boards are "
            "often informal, undisclosed arrangements centered on signaling and network access rather "
            "than governance. Fourth, the observer-to-director pipeline is active: Andreessen "
            "(Coinbase), Farrell (Reddit), and Brand (Bumble) all transitioned from observer to "
            "director, suggesting that observer seats serve as a proving ground. Fifth, observer "
            "rights almost universally terminate at IPO due to Reg FD concerns, with exceptions only "
            "for very large strategic shareholders (Advance at Reddit) \u2014 and even Microsoft\u2019s "
            "observer seat at OpenAI was ultimately relinquished under regulatory pressure."
        ).font.name = "Times New Roman"
        for run in para.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        print(f"Updated interpretive paragraph at {i}")
        break

# Save
output = "C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx"
doc.save(output)
print(f"Saved updated v3 to: {output}")
