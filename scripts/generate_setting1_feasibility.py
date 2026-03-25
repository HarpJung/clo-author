"""Generate Word document: Setting 1 Feasibility Study — S-1 Mining & Capital IQ."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15


def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h


def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = str(val)
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.size = Pt(9)
    return table


def add_bold_para(bold_text, normal_text=""):
    p = doc.add_paragraph()
    run_b = p.add_run(bold_text)
    run_b.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


# =====================================================================
# TITLE PAGE
# =====================================================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Setting 1 Feasibility Study:\nS-1 Mining & Capital IQ Database Exploration")
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Board Governance Architecture and Post-IPO Outcomes\nThree-Tier System: Directors, Observers, and Advisors")
run.font.size = Pt(13)
run.italic = True

doc.add_paragraph()
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run("Harp Jung\nHarvard University\nMarch 2026")
run.font.size = Pt(12)

doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run("Data Feasibility Assessment")
run.italic = True

doc.add_page_break()

# =====================================================================
# TABLE OF CONTENTS (manual)
# =====================================================================
add_heading_styled("Contents", level=1)
toc_items = [
    "1. Executive Summary",
    "2. S-1 Filing Analysis",
    "   2.1 EDGAR Full-Text Search Prevalence",
    "   2.2 S-1 Governance Architecture: 18-Company Sample",
    "   2.3 Key Patterns from S-1 Analysis",
    "   2.4 S-1 Data Limitations",
    "3. Capital IQ (WRDS) Database Exploration",
    "   3.1 CIQ Company Universe",
    "   3.2 Observer Coverage in CIQ",
    "   3.3 Advisory Board Coverage in CIQ",
    "   3.4 CIQ Data Structure",
    "4. BoardEx (WRDS) Database Exploration",
    "5. Other Datasets Considered",
    "6. Database Comparison Matrix",
    "7. Recommended Data Strategy",
    "8. Next Steps",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# =====================================================================
# 1. EXECUTIVE SUMMARY
# =====================================================================
add_heading_styled("1. Executive Summary", level=1)

doc.add_paragraph(
    "This document reports the results of a feasibility study for Setting 1 of the "
    "three-tier board governance research agenda: using S-1 registration statements and "
    "the Capital IQ database to study how pre-IPO governance architectures (directors, "
    "observers, and advisors) affect post-IPO outcomes."
)

doc.add_paragraph("Key findings:")

add_bullet("S-1 filings contain rich, extractable observer data. "
           "EDGAR full-text search identifies 1,133 files in S-1 filings (2020-2024) "
           "mentioning \"board observer.\" A sample of 18 VC/PE-backed IPOs found observer "
           "provisions in 14 of 17 confirmed cases (82%).",
           "S-1 Mining: ")

add_bullet("Capital IQ on WRDS tracks board observers for 2,602 private companies "
           "and advisory board members for 21,697 private companies, with named individuals, "
           "titles, and current/former flags. CIQ has board data for 345,099 private companies total.",
           "Capital IQ: ")

add_bullet("PitchBook has a dedicated \"Board Members & Observers\" section with historical "
           "data, but bulk download is difficult. It serves as a validation and enrichment source.",
           "PitchBook: ")

add_bullet("BoardEx has only 31 observer records. It is not useful for this research.",
           "BoardEx: ")

add_bullet("No existing dataset systematically tracks board observers for private firms. "
           "This data gap is itself a contribution — the governance literature studies directors "
           "because that is what the data captures, while the actual governance architecture "
           "includes two additional tiers that are largely data-invisible.",
           "The Observer Data Gap: ")

doc.add_page_break()

# =====================================================================
# 2. S-1 FILING ANALYSIS
# =====================================================================
add_heading_styled("2. S-1 Filing Analysis", level=1)

add_heading_styled("2.1 EDGAR Full-Text Search Prevalence", level=2)

doc.add_paragraph(
    "Using the SEC EDGAR Full-Text Search System (EFTS) API with proper User-Agent headers, "
    "we searched for governance-related terms across all S-1 registration statements. "
    "The EFTS API endpoint (efts.sec.gov/LATEST/search-index) returns individual file-level "
    "hits within S-1 filings, including exhibits such as Investor Rights Agreements, "
    "Stockholders Agreements, and Board Observer Agreements."
)

add_table(
    ["Search Term", "Files in S-1s (2020-2024)", "Files in S-1s (All Years)"],
    [
        ['"board observer"', "1,133", "2,749"],
        ['"observer rights"', "429", "—"],
        ['"non-voting observer"', "242", "—"],
        ['"advisory board"', "5,563", "—"],
        ['"board of advisors"', "1,066", "—"],
        ['"board advisor"', "299", "—"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "From the first 100 results for \"board observer\" (2020-2024), we identified 77 unique companies. "
    "These span VC-backed tech/biotech IPOs, PE-backed energy and healthcare companies, and SPACs."
)

add_heading_styled("2.2 S-1 Governance Architecture: 18-Company Sample", level=2)

doc.add_paragraph(
    "We analyzed S-1 filings for 18 companies using two sampling strategies:"
)
add_bullet("Convenience sample (10 companies): Well-known VC/PE-backed IPOs selected "
           "without knowledge of observer status. This provides an unbiased estimate of observer prevalence.")
add_bullet("EDGAR search sample (8 companies): Selected from confirmed \"board observer\" "
           "full-text search hits. This provides detailed examples of observer provisions.")

doc.add_paragraph(
    "For the EDGAR search sample, we directly fetched the actual S-1 exhibit text from SEC EDGAR "
    "using Python with proper User-Agent headers, and extracted observer provisions programmatically. "
    "For the convenience sample, we used web research to analyze each filing."
)

add_heading_styled("Panel A: Convenience Sample (Unbiased)", level=3)

add_table(
    ["Company", "Year", "Backing", "Directors", "Observers", "Who Gets Observer Rights", "Advisory Board"],
    [
        ["Coinbase", "2021", "VC", "8", "4 seats", "USV, a16z, DFJ, IVP", "None"],
        ["Rivian", "2021", "VC", "7", "Yes", "Ford (terminated), Amazon (retained)", "None"],
        ["DoorDash", "2020", "VC", "9", "Yes", "Per IRA, terminated at IPO", "None"],
        ["Reddit", "2024", "VC", "7", "1", "Advance Publications (post-IPO)", "None"],
        ["Bumble", "2021", "PE", "11", "1", "Blackstone (Martin Brand; 5% threshold)", "None"],
        ["Cava Group", "2023", "PE", "10", "1", "\"CGI\" entity, nonvoting", "None"],
        ["Klaviyo", "2023", "VC", "9", "None", "\u2014", "None"],
        ["Airbnb", "2020", "VC", "9", "None", "\u2014", "Stakeholder Cmte"],
        ["Rubrik", "2024", "VC", "9", "None", "\u2014", "None"],
        ["Instacart", "2023", "VC", "8", "Unconfirmed", "\u2014", "None"],
    ]
)

doc.add_paragraph()
add_bold_para("Convenience sample hit rate: ", "6 of 9 confirmed (67%) have observer provisions.")

doc.add_paragraph()
add_heading_styled("Panel B: EDGAR Search Sample (Confirmed Hits)", level=3)

add_table(
    ["Company", "Year", "Backing", "Observers", "Who", "Threshold", "Notable Feature"],
    [
        ["Waystar", "2024", "PE", "1", "CPPIB", "\u226510% common stock", "Non-transferable; post-IPO"],
        ["LandBridge", "2024", "PE", "= # directors", "Five Point Energy", "Tiered 10-40%+", "Most extreme: majority+1"],
        ["Kodiak Gas", "2023", "PE", "= # directors", "EQT/Frontier TopCo", "Tiered 10-35%+", "2 observers at 35%+"],
        ["Pulmonx", "2020", "VC", "1", "BSC", "\u226515.15M Series F-1 shares", "Share count threshold"],
        ["Blend Labs", "2021", "VC", "Per Major Inv.", "Lightspeed + others", "Major Investor status", "Multiple observers"],
        ["Sight Sciences", "2021", "VC", "1", "D1 Capital (J. Rogers)", "\u226525% Series E+F", "Named individual"],
        ["Sera Prognostics", "2021", "VC", "1", "Baker Bros. Advisors", "\u22654% post-IPO", "Fallback; no fiduciary duty"],
        ["ATAI Life Sci", "2021", "VC", "1", "Columbia", "Until Funding Threshold", "Academic institution"],
    ]
)

doc.add_paragraph()

add_heading_styled("2.3 Key Patterns from S-1 Analysis", level=2)

add_bold_para("Pattern 1: Observer architecture varies by backing type.")
add_table(
    ["Dimension", "VC-Backed", "PE-Backed"],
    [
        ["Typical observer count", "1-4 named investors", "Mirrors director designations (1:1)"],
        ["Threshold type", "Share count or % of series", "% of total common stock"],
        ["Post-IPO survival", "Usually terminates at IPO", "Often persists post-IPO"],
        ["Examples", "Coinbase (4), Blend (per Major Inv.)", "Kodiak (2 at 35%+), LandBridge (majority+1)"],
    ]
)

doc.add_paragraph()
add_bold_para("Pattern 2: Observer-to-director pipeline is documented.")
add_bullet("Marc Andreessen: Coinbase observer (a16z seat, 7+ years via Chris Dixon) \u2192 director (Dec 2020)")
add_bullet("Sarah Farrell: Reddit observer (2021-2024) \u2192 director (May 2024)")
add_bullet("Martin Brand: Bumble observer (Blackstone) \u2192 director (Aug 2024)")

add_bold_para("Pattern 3: Zero formal advisory boards in any S-1.")
doc.add_paragraph(
    "Despite 5,563 S-1 files mentioning \"advisory board,\" none of the 18 companies in our sample "
    "disclosed a governance-level advisory board. The \"advisory board\" mentions in EDGAR likely "
    "refer to Scientific Advisory Boards, Medical Advisory Boards, and similar domain-specific "
    "bodies — not governance advisory boards."
)

add_bold_para("Pattern 4: Fiduciary language varies \u2014 directly relevant to Setting 2 (NVCA 2020).")
add_bullet("Coinbase (pre-2020 IRA): Observer must act \"in a fiduciary manner\" regarding confidential information")
add_bullet("Sera Prognostics (2021): \"The BBA Observer shall not...have or be deemed to have any fiduciary or other duties\"")
add_bullet("Others: Confidentiality agreement required; no explicit fiduciary language")

add_bold_para("Pattern 5: Ownership thresholds create natural cross-sectional variation.")
add_table(
    ["Threshold Type", "Companies", "Range"],
    [
        ["% of common stock", "Waystar, Bumble, LandBridge, Kodiak", "5-40%"],
        ["% of series preferred shares", "Sight Sciences, Sera Prognostics", "4-25%"],
        ["Absolute share count", "Pulmonx", "15.15M shares"],
        ["Major Investor dollar threshold", "Blend Labs", "Dollar amount (typical $1-5M)"],
    ]
)

doc.add_paragraph()

add_heading_styled("2.4 S-1 Data Limitations", level=2)

add_bullet("Snapshot, not history: S-1 captures the governance architecture at the time of filing. "
           "It does not systematically document how observer arrangements evolved across funding rounds.")
add_bullet("IPO firms only: S-1 analysis is limited to companies that go public \u2014 a selected sample. "
           "Companies that stay private, fail, or are acquired never file an S-1.")
add_bullet("Disclosure variation: Observer arrangements are not standardized in S-1 filings. "
           "Some companies disclose detailed provisions in IRA exhibits; others mention observers "
           "only in passing in the \"Related Party Transactions\" section.")
add_bullet("Advisory boards are informal: If a company has governance-level advisors, they "
           "are unlikely to appear in the S-1. The \"advisory board\" mentions in EDGAR refer to "
           "scientific/medical advisory boards, not governance bodies.")
add_bullet("Historical traces exist: Funding round history, IRA amendment dates, director tenure dates, "
           "and board change narratives provide some historical information. "
           "The Ewens & Malenko dataset (Form D filings) tracks board composition at each round but "
           "does not capture observers.")

doc.add_page_break()

# =====================================================================
# 3. CAPITAL IQ
# =====================================================================
add_heading_styled("3. Capital IQ (WRDS) Database Exploration", level=1)

add_heading_styled("3.1 CIQ Company Universe", level=2)

doc.add_paragraph(
    "Capital IQ (accessed via WRDS PostgreSQL) contains 37.4 million entities, of which "
    "27.1 million are classified as private companies. Board/people data is available for "
    "a subset of these:"
)

add_table(
    ["Company Type", "Full Universe", "With Board Data", "Coverage"],
    [
        ["Private Company", "27,093,707", "345,099", "1.3%"],
        ["Public Company", "71,001", "63,923", "90.0%"],
        ["Private Investment Firm", "288,095", "27,112", "9.4%"],
        ["All Types", "37,426,013", "461,599", "1.2%"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "CIQ has near-complete board coverage for public companies (90%) and selective coverage for "
    "private firms (345K of 27M). The 345K private firms with board data likely skew toward "
    "larger, more visible companies (VC/PE-backed, with press coverage or SEC filings). "
    "This is the right subsample for studying three-tier governance in institutionally-backed firms."
)

add_heading_styled("3.2 Observer Coverage in CIQ", level=2)

doc.add_paragraph(
    "CIQ tracks board observers through the \"title\" field in ciqprofessional. "
    "The following title variants were found:"
)

add_table(
    ["Title", "Records"],
    [
        ["Board Observer", "2,421"],
        ["Former Board Observer", "1,346"],
        ["Observer", "497"],
        ["Former Observer", "308"],
        ["Board Observer & Director", "28"],
        ["Former Board Observer and Director", "22"],
        ["Board Observers", "21"],
        ["Board Observer & Advisor", "18"],
        ["Co-Founder & Board Observer", "13"],
        ["Non-Voting Observer", "9"],
        ["Former Non-Voting Observer", "9"],
        ["Other variants", "~200"],
        ["TOTAL (approximate)", "~5,000"],
    ]
)

doc.add_paragraph()

add_table(
    ["Metric", "Count"],
    [
        ["Unique private companies with observer records", "2,602"],
        ["Observer records at private companies", "3,757"],
        ["Observer records at public companies", "431"],
        ["Observer records at investment firms", "82"],
    ]
)

doc.add_paragraph()
doc.add_paragraph(
    "The data includes very recent companies (founded 2023-2024) such as Xaira Therapeutics, "
    "Seaport Therapeutics, and Braveheart Bio \u2014 indicating CIQ is actively collecting "
    "observer data for current VC-backed startups. The \"Former Board Observer\" category "
    "(1,346 records) enables studying observer transitions (e.g., observer-to-director pipeline)."
)

add_heading_styled("3.3 Advisory Board Coverage in CIQ", level=2)

add_table(
    ["Title", "Records"],
    [
        ["Member of Advisory Board", "55,375"],
        ["Former Member of Advisory Board", "19,517"],
        ["Member of Scientific Advisory Board", "17,469"],
        ["Former Member of Scientific Advisory Board", "8,275"],
        ["Member of Medical Advisory Board", "4,017"],
        ["Board Advisor", "1,012"],
        ["Member of Clinical Advisory Board", "2,267"],
        ["Member of Strategic Advisory Board", "1,804"],
        ["Member of Technical Advisory Board", "1,628"],
        ["Chairman of Advisory Board", "779"],
    ]
)

doc.add_paragraph()
add_bold_para("Unique private companies with advisory board records: ", "21,697")
doc.add_paragraph(
    "The advisory board data is dominated by Scientific and Medical advisory boards, consistent "
    "with the S-1 finding that governance-level advisory boards are rare while domain-specific "
    "advisory boards are common. The \"Board Advisor\" title (1,012 records) and \"Member of "
    "Strategic Advisory Board\" (1,804 records) may be closer to governance-level advisory roles."
)

add_heading_styled("3.4 CIQ Data Structure", level=2)

doc.add_paragraph("The relevant CIQ tables on WRDS (schema: ciq_pplintel) are:")

add_bullet("ciqprofessional: Links persons to companies with title, board flag, current/former flag", "")
add_bullet("ciqperson: Person demographics (name, biography)", "")
add_bullet("ciqprotoprofunction: Maps professionals to standardized function codes with start/end dates", "")
add_bullet("ciqprofunction: Function code lookup (includes boardflag, advisorflag indicators)", "")
add_bullet("ciq_common.ciqcompany: Company details (name, type, status, year founded, industry)", "")

doc.add_paragraph(
    "CIQ's standardized function codes (ciqprofunction) include four advisory board functions: "
    "Member of Advisory Board (ID 243), Chairman of Advisory Board (ID 296), "
    "Co-Chairman of Advisory Board (ID 341), and Vice Chairman of Advisory Board (ID 391). "
    "Notably, there is NO standardized function code for \"Board Observer\" \u2014 observer data "
    "exists only in the free-text \"title\" field of ciqprofessional, suggesting it is captured "
    "from company disclosures rather than systematically coded by CIQ analysts."
)

doc.add_page_break()

# =====================================================================
# 4. BOARDEX
# =====================================================================
add_heading_styled("4. BoardEx (WRDS) Database Exploration", level=1)

doc.add_paragraph(
    "BoardEx North America (boardex_na schema) was explored for observer and advisor data. "
    "BoardEx tracks 2.2 million organizations and 1.7 million people globally."
)

add_bold_para("Observer coverage: ", "31 records with rolename = 'Observer.' This is essentially "
              "zero coverage. BoardEx does not systematically track board observers.")

add_bold_para("Advisor coverage: ", "586 records with rolename = 'Advisory Board Member,' "
              "plus 266 'Advisor,' 173 'Advisory Director,' and 142 'Senior Advisor.' "
              "Modest but not comprehensive.")

add_bold_para("Company advisors table: ", "na_company_profile_advisors tracks corporate service "
              "providers (auditors, lawyers, registrars, compensation consultants) \u2014 NOT "
              "governance-level board advisors.")

add_bold_para("Conclusion: ", "BoardEx is not useful for studying board observers. Its 31 observer "
              "records appear to be incidental captures, not systematic data collection. "
              "BoardEx remains valuable for director-level analysis (535K+ director records, "
              "network data, compensation) but cannot address the observer/advisor tiers.")

doc.add_page_break()

# =====================================================================
# 5. OTHER DATASETS
# =====================================================================
add_heading_styled("5. Other Datasets Considered", level=1)

add_table(
    ["Dataset", "Directors", "Observers", "Advisors", "Private Firms", "Bulk Download", "Notes"],
    [
        ["PitchBook", "Yes", "Yes (dedicated)", "Yes", "Yes", "Hard", "Best observer coverage; manual extraction"],
        ["Ewens & Malenko", "Yes (3 types)", "No", "No", "Yes", "Yes (GitHub)", "Form D + VentureSource; no observer data"],
        ["Crunchbase", "Yes", "No", "Yes", "Yes", "API", "Self-reported; variable quality"],
        ["VentureSource (CB Insights)", "Yes", "Unknown", "Unknown", "Yes", "Unknown", "Acquired 2020; access unclear"],
        ["Form D (SEC)", "Yes", "No", "No", "Yes", "Yes (EDGAR)", "Lists directors/officers only"],
        ["LinkedIn", "Partial", "Self-reported", "Self-reported", "Yes", "Scraping", "Legal per hiQ v. LinkedIn; novel approach"],
        ["Preqin", "Unknown", "Unknown", "Unknown", "Yes (PE/VC)", "Unknown", "No evidence of board-level data"],
        ["Delaware Corp. Filings", "No", "No", "No", "Yes", "Yes", "Charters, not IRAs"],
        ["NVCA Surveys", "Aggregate", "Aggregate (82%)", "No", "Aggregate", "No", "Prevalence only, not firm-level"],
    ]
)

doc.add_page_break()

# =====================================================================
# 6. DATABASE COMPARISON
# =====================================================================
add_heading_styled("6. Database Comparison Matrix", level=1)

add_table(
    ["Feature", "S-1 Filings", "Capital IQ (WRDS)", "PitchBook", "BoardEx (WRDS)"],
    [
        ["Observer data", "Yes (detailed)", "Yes (2,602 firms)", "Yes (dedicated)", "No (31 records)"],
        ["Advisory board data", "Rare", "Yes (21,697 firms)", "Yes", "Limited (586)"],
        ["Director data", "Yes", "Yes (345K private)", "Yes", "Yes (public focus)"],
        ["Private firm coverage", "At IPO only", "345K firms", "Broad", "Limited"],
        ["Governance provisions", "Full text (thresholds, fiduciary language)", "Title only", "Unknown", "Role name only"],
        ["Historical data", "Snapshot + traces", "Current/former flag", "Yes (historical)", "Start/end dates"],
        ["Bulk download", "Yes (EDGAR API)", "Yes (SQL)", "Hard", "Yes (SQL)"],
        ["Named individuals", "In exhibits", "Yes", "Yes", "Yes"],
        ["Ownership thresholds", "Yes (in IRA text)", "No", "Unknown", "No"],
        ["Fiduciary language", "Yes (in IRA text)", "No", "No", "No"],
    ]
)

doc.add_page_break()

# =====================================================================
# 7. RECOMMENDED DATA STRATEGY
# =====================================================================
add_heading_styled("7. Recommended Data Strategy", level=1)

doc.add_paragraph("The three data sources are complementary and serve different purposes:")

add_bold_para("Layer 1: Capital IQ (WRDS) \u2014 Broad Cross-Section")
add_bullet("Use for: Identifying the universe of firms with board observers; extensive margin analysis")
add_bullet("Extract: All ~5,000 observer records with company IDs, person IDs, titles, current/former flags")
add_bullet("Link to: CIQ company data (type, status, industry, founding year) and CIQ financials")
add_bullet("Limitation: Title field only \u2014 no governance provision details")

add_bold_para("Layer 2: S-1 Filings (EDGAR) \u2014 Deep Governance Detail for IPO Firms")
add_bullet("Use for: Intensive margin analysis; coding governance architecture in detail")
add_bullet("Extract: Observer count, identity, ownership thresholds, fiduciary language, "
           "termination conditions, post-IPO survival")
add_bullet("Method: EDGAR EFTS search to identify S-1s with observer language, then fetch exhibits "
           "with proper User-Agent headers and extract provisions via NLP/regex")
add_bullet("Limitation: IPO firms only; snapshot at filing date")

add_bold_para("Layer 3: PitchBook \u2014 Validation & Historical Enrichment")
add_bullet("Use for: Validating CIQ observer data; adding historical dimension")
add_bullet("Advantage: Dedicated \"Board Members & Observers\" section with historical tracking")
add_bullet("Limitation: Difficult to bulk extract; best used for targeted validation of CIQ/S-1 findings")

doc.add_paragraph()
doc.add_paragraph(
    "For the class write-up, present all three layers as the proposed data strategy. "
    "For the journal paper, begin with the CIQ extract (broadest coverage, easiest to obtain) "
    "and supplement with S-1 exhibit coding for the IPO subsample where deep governance "
    "detail is available."
)

doc.add_page_break()

# =====================================================================
# 8. NEXT STEPS
# =====================================================================
add_heading_styled("8. Next Steps", level=1)

add_bullet("Pull full CIQ observer extract: All records with title matching observer variants, "
           "with company details, person details, and start/end dates", "1. ")
add_bullet("Pull full CIQ advisory board extract: Same for advisory board title variants", "2. ")
add_bullet("Cross-reference CIQ observer companies with S-1 filers: "
           "Identify which CIQ companies have S-1 filings available on EDGAR for deeper analysis", "3. ")
add_bullet("PitchBook validation: For a random subsample of 50-100 CIQ observer companies, "
           "manually check PitchBook for consistency and additional historical detail", "4. ")
add_bullet("Build the S-1 governance coding protocol: Define the variables to extract from "
           "each S-1 exhibit (observer count, identity, threshold, fiduciary language, etc.) "
           "and pilot on the 18 companies already analyzed", "5. ")
add_bullet("Merge with outcome data: Link CIQ company IDs and S-1 CIKs to post-IPO outcome "
           "databases (CRSP, Compustat, Audit Analytics, ExecuComp)", "6. ")

doc.add_paragraph()

# =====================================================================
# APPENDIX: WRDS CONNECTION DETAILS
# =====================================================================
add_heading_styled("Appendix: WRDS Database Access Details", level=1)

doc.add_paragraph("Connection: wrds-pgdata.wharton.upenn.edu:9737 (PostgreSQL)")
doc.add_paragraph()

add_bold_para("Key CIQ tables for observer/advisor research:")
add_bullet("ciq_pplintel.ciqprofessional \u2014 Person-company links with titles and board flags")
add_bullet("ciq_pplintel.ciqperson \u2014 Person demographics")
add_bullet("ciq_pplintel.ciqprotoprofunction \u2014 Standardized function codes with dates")
add_bullet("ciq_pplintel.ciqprofunction \u2014 Function code lookup (boardflag, advisorflag)")
add_bullet("ciq_common.ciqcompany \u2014 Company details (type, status, founding year)")
add_bullet("ciq_common.ciqcompanytype \u2014 Company type lookup (ID 5 = Private Company)")
add_bullet("ciq_common.ciqcountrygeo \u2014 Country geography")

add_bold_para("Key BoardEx tables (limited utility):")
add_bullet("boardex_na.na_wrds_org_composition \u2014 Role-level data (rolename field)")
add_bullet("boardex_na.na_board_unlisted_assoc \u2014 Associations with private companies")

add_bold_para("Key EDGAR access:")
add_bullet("EFTS API: efts.sec.gov/LATEST/search-index?q=\"board+observer\"&forms=S-1")
add_bullet("Requires User-Agent header (e.g., 'ResearcherName email@university.edu')")
add_bullet("S-1 exhibits fetchable via www.sec.gov/Archives/edgar/data/{CIK}/{accession}/{filename}")

# =====================================================================
# SAVE
# =====================================================================
output_path = os.path.normpath(os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "..",
    "Setting 1 Feasibility -- S-1 Mining & Capital IQ.docx"
))
doc.save(output_path)
print(f"Saved to: {output_path}")
