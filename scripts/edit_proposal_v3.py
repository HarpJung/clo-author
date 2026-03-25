"""Edit Research Proposal v2 -> v3: Add Section 1.0 Definitions."""

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from copy import deepcopy

# Open v2
doc = docx.Document("C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v2 -- Three-Tier Board Governance.docx")

# Find the paragraph index for "Section 1: What Are Board Observers?"
section1_idx = None
for i, para in enumerate(doc.paragraphs):
    if "Section 1" in para.text and "What Are Board Observers" in para.text:
        section1_idx = i
        break

if section1_idx is None:
    print("ERROR: Could not find Section 1 heading")
    exit()

print(f"Found Section 1 at paragraph {section1_idx}")

# We need to insert new paragraphs AFTER the Section 1 heading (paragraph index section1_idx)
# and BEFORE paragraph section1_idx+1 (the current intro text)
# In python-docx, we can insert by manipulating the XML directly

from docx.oxml.ns import qn
from lxml import etree

body = doc.element.body

# Get the XML element for the Section 1 heading
section1_elem = doc.paragraphs[section1_idx]._element

# Build all the new content as a list of paragraphs to insert
new_content = []

# Helper to create a paragraph with specific formatting
def make_para(text, bold=False, italic=False, size=12, indent_left=0, indent_right=0, style_name="Normal"):
    p = docx.oxml.OxmlElement("w:p")
    # Set paragraph properties
    pPr = docx.oxml.OxmlElement("w:pPr")
    # Style
    pStyle = docx.oxml.OxmlElement("w:pStyle")
    pStyle.set(qn("w:val"), style_name)
    pPr.append(pStyle)
    # Spacing (double)
    spacing = docx.oxml.OxmlElement("w:spacing")
    spacing.set(qn("w:line"), "480")
    spacing.set(qn("w:lineRule"), "auto")
    pPr.append(spacing)
    # Indentation
    if indent_left > 0 or indent_right > 0:
        ind = docx.oxml.OxmlElement("w:ind")
        if indent_left > 0:
            ind.set(qn("w:left"), str(int(indent_left * 1440)))  # inches to twips
        if indent_right > 0:
            ind.set(qn("w:right"), str(int(indent_right * 1440)))
        pPr.append(ind)
    p.append(pPr)
    # Run
    r = docx.oxml.OxmlElement("w:r")
    rPr = docx.oxml.OxmlElement("w:rPr")
    # Font
    rFonts = docx.oxml.OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), "Times New Roman")
    rFonts.set(qn("w:hAnsi"), "Times New Roman")
    rPr.append(rFonts)
    # Size
    sz = docx.oxml.OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    rPr.append(sz)
    if bold:
        b = docx.oxml.OxmlElement("w:b")
        rPr.append(b)
    if italic:
        i_elem = docx.oxml.OxmlElement("w:i")
        rPr.append(i_elem)
    r.append(rPr)
    t = docx.oxml.OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    r.append(t)
    p.append(r)
    return p

# Actually, manipulating XML directly is complex. Let me use a simpler approach:
# Insert paragraphs using the document's add_paragraph method at specific positions.
# python-docx doesn't have insert_paragraph_before, but we can use the XML approach.

# Simpler approach: rebuild the document by copying everything and inserting new content

# Let's use a different strategy - insert after the section 1 heading element
# by adding new elements after it in the XML tree

def insert_after(parent, ref_elem, new_elem):
    """Insert new_elem after ref_elem in parent."""
    ref_idx = list(parent).index(ref_elem)
    parent.insert(ref_idx + 1, new_elem)

# We'll insert in reverse order (last first) so each goes right after the section heading
# Actually let's collect all elements and insert them in order

# First, let's just use the approach of adding to the end and relying on position
# Better approach: create the heading and paragraphs as actual docx paragraph objects

# The cleanest approach: save v2, create a new document for v3 by copying with insertions
# But that's complex too.

# Simplest working approach: use python-docx's internal XML manipulation
# Insert new paragraph elements after the Section 1 heading

insert_point = section1_elem  # insert after this

# Create heading "1.0 Definitions: The Three Governance Roles"
h2 = doc.add_heading("1.0 Definitions: The Three Governance Roles", level=2)
h2_elem = h2._element
body.remove(h2_elem)  # remove from end
insert_point.addnext(h2_elem)
insert_point = h2_elem

# Now add paragraphs after the heading
texts_to_add = [
    ("normal", "Before examining how board observers function in practice, it is essential to define the three distinct governance roles that constitute the board-level architecture of venture-backed firms. While these roles are sometimes conflated in popular discussion, they differ fundamentally in their legal authority, fiduciary obligations, information access, and accountability structures."),

    ("bold_normal", "Board of Directors (Tier 1).",
     ' Directors are formally elected members of the corporation\'s governing body who hold voting rights, bear fiduciary duties of care and loyalty, and face personal liability for their governance decisions. In the VC context, Ewens and Malenko (2024) classify directors into three subtypes: executive directors (founders/officers of the startup), VC directors (representatives of the investing VC firm), and independent directors (outsiders unaffiliated with either the entrepreneur or the investors). Their panel of 7,780 startups shows the average board has 4.5 directors: "approximately 2 seats held by VCs, 1.7 by executives, and 0.8 by independent directors" (Ewens & Malenko, 2024, p. 4). Directors\' fiduciary duties require them to "maximiz[e] the value of the firm as a whole" (p. 5), and they can be held personally liable for breaching these duties.'),

    ("bold_normal", "Board Observers (Tier 2).",
     ' Board observers are nonvoting attendees who receive full board materials, attend all meetings, and participate in discussions, but hold no voting rights and owe no fiduciary duties to the company or its shareholders. Packin and Alon-Beck (2025, p. 1508) define them as "individuals without voting rights who serve as strategic eyes on the board, providing insight and influence without the fiduciary responsibilities of board membership. Their role exemplifies control without formal control." Observers are "typically appointed through contractual instruments such as investment agreements, term sheets, or side letters\u2014though rarely mentioned in bylaws or articles of association" (p. 1535). Ewens and Malenko (2024) acknowledge observers as a distinct governance mechanism in their footnote 3, noting that rather than giving advisors voting power, "the firm could establish an advisory board (as is frequently done in private firms) or add these individuals to the board as \'board observers\'" (p. 5). The NVCA Q4 2023 CFO Survey documents that 82% of VC entities employ board observers, with 100% adoption among entities with AUM exceeding $500 million (Packin & Alon-Beck, 2025, p. 1514\u20131534).'),

    ("bold_normal", "Board Advisors (Tier 3).",
     ' Advisory board members occupy the most informal tier of the governance architecture. They typically serve in an honorary capacity, providing domain expertise, network access, or reputational signaling without any formal governance role. Unlike observers, advisors generally do not attend board meetings, do not receive board materials, and have no contractual information rights. Their engagement is episodic rather than systematic. Ewens and Malenko (2024, fn. 3) reference advisory boards as a common structure in private firms, distinct from both voting directors and board observers. Pollman (2024) situates advisory roles within the broader context of governance complexity, noting that "startup governance typically evolves from relatively simple to very complex sets of tensions between and among participants" (p. 13), with advisory roles emerging as one mechanism to manage this complexity without expanding the formal board.'),

    ("normal", "The following table summarizes the key distinctions across these three tiers:"),
]

for item in texts_to_add:
    if item[0] == "normal":
        p = doc.add_paragraph(item[1])
        p.style = doc.styles["Normal"]
        for run in p.runs:
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(6)
        p_elem = p._element
        body.remove(p_elem)
        insert_point.addnext(p_elem)
        insert_point = p_elem

    elif item[0] == "bold_normal":
        p = doc.add_paragraph()
        p.style = doc.styles["Normal"]
        run_b = p.add_run(item[1])
        run_b.bold = True
        run_b.font.name = "Times New Roman"
        run_b.font.size = Pt(12)
        run_n = p.add_run(item[2])
        run_n.font.name = "Times New Roman"
        run_n.font.size = Pt(12)
        p.paragraph_format.line_spacing = 2.0
        p.paragraph_format.space_after = Pt(6)
        p_elem = p._element
        body.remove(p_elem)
        insert_point.addnext(p_elem)
        insert_point = p_elem

# Now add the comparison table
table = doc.add_table(rows=8, cols=4)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Headers
headers = ["Attribute", "Board of Directors (Tier 1)", "Board Observers (Tier 2)", "Board Advisors (Tier 3)"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

# Data rows
rows_data = [
    ["Voting Rights", "Yes \u2014 full voting power on board decisions", "No \u2014 nonvoting attendee", "No \u2014 no board role"],
    ["Fiduciary Duty", "Yes \u2014 duty of care and loyalty to all shareholders", "No \u2014 no fiduciary obligation (Packin & Alon-Beck, 2025; Obasi v. Tibet, 3d Cir. 2019)", "No \u2014 no legal obligation (Toj interview)"],
    ["Information Access", "Full \u2014 all board materials, discussions, executive sessions", "Full \u2014 same materials as directors, attends all meetings (may be excluded for privilege/trade secrets)", "Limited \u2014 no regular access to board materials; ad hoc consultation"],
    ["Personal Liability", "Yes \u2014 mitigated by D&O insurance and business judgment rule", "No \u2014 not liable unless acting as de facto director", "No \u2014 no governance liability"],
    ["Legal Basis", "Corporate charter, bylaws, shareholder vote", "Investor Rights Agreement, Stockholders Agreement, or side letter (contractual)", "Informal agreement, sometimes equity compensation"],
    ["Appointment", "Elected by shareholders or designated per charter", "Designated by specific investors per IRA provisions", "Invited by management; personal/professional network"],
    ["Prevalence in VC", "Universal \u2014 required by law", "82% of VC entities use them (NVCA 2024); 100% of entities >$500M AUM", "Common but informal; not systematically tracked"],
]

for r_idx, row_data in enumerate(rows_data):
    for c_idx, cell_text in enumerate(row_data):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(9)
                run.font.name = "Times New Roman"

# Move table element after insert point
table_elem = table._element
body.remove(table_elem)
insert_point.addnext(table_elem)
insert_point = table_elem

# Add a source note after the table
note = doc.add_paragraph()
note.style = doc.styles["Normal"]
run = note.add_run("Sources: Director definitions from Ewens & Malenko (2024); Observer definitions from Packin & Alon-Beck (2025, pp. 1508, 1535, 1537); NVCA survey statistics from Packin & Alon-Beck (2025, pp. 1514\u20131535); Advisor characterization from practitioner interviews (Matt, Toj) and Ewens & Malenko (2024, fn. 3).")
run.font.name = "Times New Roman"
run.font.size = Pt(10)
run.italic = True
note.paragraph_format.space_after = Pt(12)
note_elem = note._element
body.remove(note_elem)
insert_point.addnext(note_elem)
insert_point = note_elem

# Add transition paragraph
transition = doc.add_paragraph(
    "With these definitions established, we now turn to examining how these roles function in practice, "
    "beginning with the critical observation that the formal distinctions outlined above mask a more "
    "complex reality in which observers wield substantial informal influence."
)
transition.style = doc.styles["Normal"]
for run in transition.runs:
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
transition.paragraph_format.line_spacing = 2.0
transition.paragraph_format.space_after = Pt(6)
trans_elem = transition._element
body.remove(trans_elem)
insert_point.addnext(trans_elem)

# Renumber existing subsections: 1.1 -> 1.1 (stays same since we added 1.0)
# The existing subsections are already 1.1, 1.2, etc. so they're fine.

# Save as v3
output = "C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx"
doc.save(output)
print(f"Saved v3 to: {output}")
