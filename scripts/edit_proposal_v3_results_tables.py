"""Edit v3: Add results tables to each Section 4 subsection."""

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = docx.Document("C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx")

body = doc.element.body

def find_para(search_text):
    for i, para in enumerate(doc.paragraphs):
        if search_text in para.text:
            return i
    return None

def insert_table_after(para_idx, headers, data, font_size=9, title=None):
    """Insert a table after the given paragraph index."""
    ref_elem = doc.paragraphs[para_idx]._element

    # Add title if provided
    if title:
        tp = doc.add_paragraph()
        tp.style = doc.styles["Normal"]
        run = tp.add_run(title)
        run.bold = True
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        tp.paragraph_format.space_before = Pt(12)
        tp.paragraph_format.space_after = Pt(4)
        tp_elem = tp._element
        body.remove(tp_elem)
        ref_elem.addnext(tp_elem)
        ref_elem = tp_elem

    # Create table
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.name = "Times New Roman"

    for r_idx, row_data in enumerate(data):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Times New Roman"

    table_elem = table._element
    body.remove(table_elem)
    ref_elem.addnext(table_elem)
    return table_elem


# =====================================================================
# 4.1 Test 1: Observer Presence
# =====================================================================
# Find the last paragraph before 4.2
idx_42 = find_para("4.2 Test 2: Observer Intensity")
# Insert before 4.2 heading = after the last paragraph of 4.1
# Find the paragraph just before 4.2
insert_idx = idx_42 - 1
while not doc.paragraphs[insert_idx].text.strip():
    insert_idx -= 1

print(f"Inserting Test 1 table after paragraph {insert_idx}")

insert_table_after(insert_idx,
    ["Outcome", "Treatment", "Control", "Coef", "t-stat", "p-value", "Sig", "N", "Controls"],
    [
        ["Analyst coverage", "Observer S-1", "Non-observer S-1", "+0.96", "2.84", "0.005", "***", "605", "Size + Lev + Year FE + Ind FE"],
        ["IPO underpricing", "Observer S-1", "Non-observer S-1", "\u22120.024", "\u22122.27", "0.023", "**", "271", "Size + Lev + Year FE"],
        ["Return volatility", "Observer S-1", "Non-observer S-1", "\u22120.066", "\u22122.05", "0.041", "**", "1,520", "Size + Lev + Year FE + Ind FE"],
        ["6-month BHAR", "Observer S-1", "Non-observer S-1", "+0.027", "0.91", "0.361", "", "1,517", "Size + Lev + Year FE"],
        ["Forecast dispersion", "Observer S-1", "Non-observer S-1", "+0.011", "0.19", "0.850", "", "559", "Size + Lev + Year FE"],
    ],
    font_size=8,
    title="Table 4.1: Test 1 Results \u2014 Observer Presence and Post-IPO Outcomes"
)

# =====================================================================
# 4.2 Test 2: Observer Intensity
# =====================================================================
idx_43 = find_para("4.3 Test 5: Full CIQ Private Firm")
insert_idx = idx_43 - 1
while not doc.paragraphs[insert_idx].text.strip():
    insert_idx -= 1

print(f"Inserting Test 2 table after paragraph {insert_idx}")

insert_table_after(insert_idx,
    ["Outcome", "Treatment", "Coef", "t / z", "p-value", "Sig", "N"],
    [
        ["Return volatility", "High observer (binary)", "\u22120.163", "\u22121.97", "0.049", "**", "184"],
        ["Return volatility", "Observer ratio (cont.)", "\u22120.951", "\u22121.70", "0.089", "*", "184"],
    ],
    font_size=9,
    title="Table 4.2: Test 2 Results \u2014 Observer Intensity and Post-IPO Volatility (CIQ \u2192 CRSP matched)"
)

# =====================================================================
# 4.3 Test 5: Full CIQ Private Firms
# =====================================================================
idx_44 = find_para("4.4 Test 3: Information Spillover")
insert_idx = idx_44 - 1
while not doc.paragraphs[insert_idx].text.strip():
    insert_idx -= 1

print(f"Inserting Test 5 table after paragraph {insert_idx}")

insert_table_after(insert_idx,
    ["Outcome", "Treatment", "OR", "Coef", "p-value", "Sig", "N", "Controls"],
    [
        ["Has lawsuit", "High observer", "0.46", "\u22120.77", "<0.001", "***", "2,537", "Board + Adv + US + Capital + Age"],
        ["Was acquired", "High observer", "0.51", "\u22120.67", "0.002", "***", "2,537", "Board + Adv + US + Capital + Age"],
        ["Has restatement", "High observer", "0.25", "\u22121.39", "0.216", "", "2,537", "Board + Adv + US + Capital + Age"],
        ["Has bankruptcy", "High observer", "0.70", "\u22120.36", "0.407", "", "2,537", "Board + Adv + US + Capital + Age"],
        ["Company failed", "High observer", "1.06", "+0.05", "0.859", "", "2,537", "Board + Adv + US + Capital + Age"],
        ["", "", "", "", "", "", "", ""],
        ["N exec changes (OLS)", "High observer", "\u2014", "\u22121.35", "<0.001", "***", "2,537", "Board + Adv + US + Capital + Age"],
        ["N lawsuits (OLS)", "High observer", "\u2014", "\u22120.12", "<0.001", "***", "2,537", "Board + Adv + US + Capital + Age"],
        ["", "", "", "", "", "", "", ""],
        ["Has lawsuit", "Observer ratio", "0.03", "\u22123.50", "0.001", "***", "2,537", "Board + Adv + US + Capital + Age"],
        ["Was acquired", "Observer ratio", "0.11", "\u22122.18", "0.003", "***", "2,537", "Board + Adv + US + Capital + Age"],
        ["Has restatement", "Observer ratio", "0.00", "\u221217.73", "0.039", "**", "2,537", "Board + Adv + US + Capital + Age"],
    ],
    font_size=8,
    title="Table 4.3: Test 5 Results \u2014 Full CIQ Private Firm Outcomes (Logistic + OLS)"
)

# =====================================================================
# 4.4 Test 3: Information Spillover
# =====================================================================
idx_45 = find_para("4.5 Test 4: Pre-Announcement")
insert_idx = idx_45 - 1
while not doc.paragraphs[insert_idx].text.strip():
    insert_idx -= 1

print(f"Inserting Test 3 table after paragraph {insert_idx}")

insert_table_after(insert_idx,
    ["Comparison", "CAR[-5,+5]", "N", "t-stat", "p-value", "Sig", "Clustering"],
    [
        ["Overall spillover (intercept)", "+0.16%", "70,218", "2.24", "0.025", "**", "VC firm (1,353)"],
        ["Same-industry pairs", "+0.58%", "24,222", "4.64", "<0.001", "***", "Unclustered"],
        ["Different-industry pairs", "\u22120.06%", "45,996", "\u22120.98", "0.329", "", "Unclustered"],
        ["Same vs. different (diff)", "+0.65%", "\u2014", "4.56", "<0.001", "***", "Unclustered"],
        ["", "", "", "", "", "", ""],
        ["Regression: same_industry", "+0.52%", "70,218", "2.21", "0.027", "**", "VC firm (1,353)"],
        ["Regression: is_director", "\u22120.34%", "70,218", "\u22122.41", "0.016", "**", "VC firm (1,353)"],
        ["", "", "", "", "", "", ""],
        ["Non-director connections", "+0.31%", "20,217", "3.53", "<0.001", "***", "Unclustered"],
        ["Director connections", "+0.10%", "50,001", "1.27", "0.203", "", "Unclustered"],
    ],
    font_size=8,
    title="Table 4.4: Test 3 Results \u2014 Cross-Portfolio Return Spillovers (Daily CAR[-5,+5])"
)

# =====================================================================
# 4.5 Test 4: Pre-Announcement Drift
# =====================================================================
idx_46 = find_para("4.6 Natural Experiments")
insert_idx = idx_46 - 1
while not doc.paragraphs[insert_idx].text.strip():
    insert_idx -= 1

print(f"Inserting Test 4 table after paragraph {insert_idx}")

insert_table_after(insert_idx,
    ["Measure", "Observer Firms", "Non-Observer Firms", "Difference", "p-value", "Sig", "N"],
    [
        ["Signed CAR[-10,-1]", "+0.47%", "+0.42%", "+0.05%", "0.907", "", "37,710"],
        ["|CAR[-10,-1]| (absolute)", "10.1%", "10.8%", "\u22120.75%", "0.016", "**", "37,710"],
        ["|CAR| w/ Year + Ind FE", "\u2014", "\u2014", "\u22121.42%", "<0.001", "***", "37,710"],
        ["", "", "", "", "", "", ""],
        ["Earnings drift (signed)", "+0.93%", "+0.55%", "+0.38%", "0.033", "**", "29,231"],
        ["Exec change drift (signed)", "\u22120.99%", "\u22120.04%", "\u22120.96%", "0.220", "", "8,479"],
    ],
    font_size=8,
    title="Table 4.5: Test 4 Results \u2014 Pre-Announcement Drift at Observer vs. Non-Observer Firms"
)

# =====================================================================
# Add a master summary table at the end of Section 4
# =====================================================================
idx_sec5 = find_para("Section 5: Interview Evidence")
insert_idx = idx_sec5 - 1
while not doc.paragraphs[insert_idx].text.strip():
    insert_idx -= 1

print(f"Inserting master summary table after paragraph {insert_idx}")

insert_table_after(insert_idx,
    ["Test", "Research Question", "Key Finding", "Coef", "p-value", "N", "Causal?"],
    [
        ["1", "Do observer firms differ post-IPO?", "+1 analyst coverage", "+0.96", "0.005***", "605", "Cross-sectional"],
        ["2", "Does observer intensity matter?", "Lower volatility", "\u22120.163", "0.049**", "184", "Cross-sectional"],
        ["5", "Private firm outcomes?", "48% fewer lawsuits", "OR=0.46", "<0.001***", "2,537", "Cross-sectional"],
        ["5", "", "49% fewer acquisitions", "OR=0.51", "0.002***", "2,537", ""],
        ["5", "", "Fewer exec changes", "\u22121.35", "<0.001***", "2,537", ""],
        ["3", "Info flows through network?", "Same-industry spillover", "+0.58%", "<0.001***", "70,218", "Quasi-causal"],
        ["3", "Accountability constrains flow?", "Director coef negative", "\u22120.34%", "0.016**", "70,218", "Quasi-causal"],
        ["4", "Insider trading via observers?", "LESS pre-announcement drift", "\u22121.42%", "<0.001***", "37,710", "Quasi-causal"],
    ],
    font_size=8,
    title="Table 4.6: Master Summary of All Preliminary Results"
)

# Save
output = "C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx"
doc.save(output)
print(f"\nSaved updated v3 to: {output}")
