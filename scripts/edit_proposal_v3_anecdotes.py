"""Edit v3: Insert Section 1.5 anecdotes before the Accountability Gap section."""

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

doc = docx.Document("C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx")

# Find "1.5 The Accountability Gap" heading
target_idx = None
for i, para in enumerate(doc.paragraphs):
    if "1.5" in para.text and "Accountability Gap" in para.text:
        target_idx = i
        break

if target_idx is None:
    print("ERROR: Could not find 1.5 Accountability Gap heading")
    exit()

print(f"Found 1.5 Accountability Gap at paragraph {target_idx}")

# We'll renumber: current 1.5 becomes 1.6
# Insert new 1.5 before it
old_heading = doc.paragraphs[target_idx]

# Change "1.5" to "1.6" in the existing heading
if "1.5" in old_heading.text:
    for run in old_heading.runs:
        if "1.5" in run.text:
            run.text = run.text.replace("1.5", "1.6")
    print("Renumbered Accountability Gap to 1.6")

body = doc.element.body
insert_point = old_heading._element

# Helper to add paragraph before insert_point
def add_before(text, bold=False, italic=False, font_size=12):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(6)
    elem = p._element
    body.remove(elem)
    insert_point.addprevious(elem)
    return elem

def add_heading_before(text, level=2):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = "Times New Roman"
    elem = h._element
    body.remove(elem)
    insert_point.addprevious(elem)
    return elem

def add_bold_normal_before(bold_text, normal_text):
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    rb = p.add_run(bold_text)
    rb.bold = True
    rb.font.name = "Times New Roman"
    rb.font.size = Pt(12)
    rn = p.add_run(normal_text)
    rn.font.name = "Times New Roman"
    rn.font.size = Pt(12)
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(6)
    elem = p._element
    body.remove(elem)
    insert_point.addprevious(elem)
    return elem

# Insert new Section 1.5
add_heading_before("1.5 The Three Tiers in Practice: Five Company Examples", level=2)

add_before(
    "To illustrate how the three-tier governance architecture manifests in practice, we present "
    "five VC-backed companies from our S-1 sample. Each example demonstrates the typical "
    "composition of each tier: founders and key investors as directors, VC and strategic investors "
    "as observers, and prominent industry figures as advisors. The pattern is remarkably consistent "
    "across sectors and company stages."
)

# Company 1: Coinbase
add_bold_normal_before(
    "Coinbase (IPO 2021, Cryptocurrency Exchange). ",
    "Coinbase\u2019s pre-IPO board had 8 directors: co-founder and CEO Brian Armstrong, "
    "co-founder Frederick Ehrsam (previously a trader at Goldman Sachs), and six outside directors "
    "including Kathryn Haun (former federal prosecutor and Stanford Law lecturer), Kelly Kramer "
    "(former CFO of Cisco Systems), Gokul Rajaram (a technology executive then at DoorDash), and "
    "Fred Wilson (co-founder of Union Square Ventures, one of the most prominent VC firms in New York). The IRA granted nonvoting observer seats to "
    "four VC investors: Union Square Ventures (Series B), Andreessen Horowitz (Series C), "
    "DFJ (Series C), and IVP (Series D). Notably, Marc Andreessen served as a16z\u2019s board observer "
    "for seven years via Chris Dixon before transitioning to a full director seat in December 2020 \u2014 "
    "illustrating the observer-to-director pipeline. Post-IPO, Coinbase established a separate "
    "Global Advisory Council featuring former SEC Chair Jay Clayton, former Senator Patrick Toomey, "
    "and former Fed President Bill Dudley \u2014 high-profile figures providing regulatory and policy "
    "expertise without any governance role."
)

# Company 2: Reddit
add_bold_normal_before(
    "Reddit (IPO 2024, Social Media). ",
    "Reddit\u2019s board had 7 directors at IPO, including co-founder and CEO Steven Huffman, Chairman "
    "David Habiger, and Vice Chairman Robert Sauerberg (an Advance Publications designee). "
    "Advance Publications (~30% owner) negotiated both 2 director designees AND 1 nonvoting observer "
    "seat via a Governance Agreement \u2014 a rare case where the same investor holds board seats and an "
    "observer seat simultaneously. Sarah Farrell served as Advance\u2019s observer from 2021 to 2024, "
    "then was elevated to full director in May 2024 \u2014 another observer-to-director transition. "
    "Reddit\u2019s observer provisions explicitly state the observer \u201cshall not\u2026have or be deemed to "
    "have, or otherwise be subject to, any duties (fiduciary or otherwise) to the Company or its "
    "stockholders.\u201d"
)

# Company 3: Bumble
add_bold_normal_before(
    "Bumble (IPO 2021, Dating App \u2014 PE-Backed). ",
    "Bumble\u2019s board had 11 directors \u2014 notably, eight of whom were women. Directors included "
    "founder Whitney Wolfe Herd (CEO, previously co-founder of Tinder), Ann Mather (Chair, a former "
    "Pixar CFO who also serves on the boards of Alphabet, Airbnb, and Netflix), and Blackstone "
    "designees Sachin Bavishi and Jonathan Korngold (both senior Blackstone investment professionals). Per the Stockholders Agreement, Blackstone designated Martin Brand \u2014 Senior Managing "
    "Director and Co-Head of U.S. Acquisitions \u2014 as a nonvoting board observer, with the right to "
    "\u201creceive at the same time and in the same manner as the Directors copies of all materials "
    "given to Directors.\u201d Brand was later elevated from observer to full director in August 2024. "
    "This PE example demonstrates that observers in PE-backed companies tend to be senior investment "
    "professionals who maintain a parallel information channel alongside their firm\u2019s director seats."
)

# Company 4: Rivian
add_bold_normal_before(
    "Rivian (IPO 2021, Electric Vehicles). ",
    "Rivian\u2019s 7-member board at IPO included founder and CEO R.J. Scaringe (an MIT-trained "
    "mechanical engineer who founded Rivian in 2009) and Amazon\u2019s Peter Krawiec (Amazon\u2019s "
    "Senior Vice President of Worldwide Corporate and Business Development, appointed under a "
    "Director Nomination Agreement). The IRA granted observer rights "
    "to both Amazon and Ford Motor Company as major strategic investors. Ford\u2019s board involvement "
    "illustrates the lifecycle dynamic: Joe Hinrichs served as Ford\u2019s representative until 2020, "
    "followed by Alexandra Ford English, then Doug Power \u2014 all of whom exited before the IPO as "
    "Ford stated: \u201cWe had a board seat as Rivian was a private company, but as they go public, "
    "Ford won\u2019t have a presence on the board.\u201d Amazon retained both a director seat AND observer "
    "rights post-IPO \u2014 reflecting its 20% ownership and strategic partnership for 100,000 electric "
    "delivery vans."
)

# Company 5: DoorDash
add_bold_normal_before(
    "DoorDash (IPO 2020, Food Delivery). ",
    "DoorDash\u2019s 9-member board included all three co-founders (Tony Xu as CEO and Chairman, Andy Fang "
    "as CTO, Stanley Tang as Chief Product Officer), venture directors John Doerr (legendary Kleiner "
    "Perkins partner who also backed Google and Amazon), Alfred Lin (Sequoia Capital partner and "
    "former Chairman/COO of Zappos), and Greg Mondre (Silver Lake co-CEO), plus independent directors "
    "Shona Brown (former SVP at Google), Andy Conrad (founder of Verily/Google Life Sciences), and "
    "William Gurley (general partner at Benchmark Capital). The Investors\u2019 Rights Agreement granted observer rights "
    "to additional VC investors not represented by a director seat \u2014 the standard arrangement where "
    "the lead investor gets a board seat and co-investors receive observer rights. All observer "
    "rights terminated at IPO, consistent with the Reg FD-driven pattern we document across our "
    "S-1 sample."
)

# Summary table
add_before(
    "The following table summarizes the three-tier architecture across these five companies:"
)

# Create summary table
table = doc.add_table(rows=6, cols=4)
table.style = "Light Grid Accent 1"
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ["Company", "Directors (Tier 1)", "Observers (Tier 2)", "Advisors (Tier 3)"]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for p in cell.paragraphs:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(9)
            run.font.name = "Times New Roman"

data = [
    ["Coinbase\n(2021)",
     "8 directors: 2 co-founders, 4 VC-affiliated (Kleiner, Sequoia, a16z, USV), 2 independent (former prosecutor, former Cisco CFO)",
     "4 observers: USV (Series B), a16z (Series C), DFJ (Series C), IVP (Series D). All terminated at IPO.",
     "Global Advisory Council (est. 2023): former SEC Chair Jay Clayton, former Sen. Pat Toomey, former Fed Pres. Bill Dudley"],
    ["Reddit\n(2024)",
     "7 directors: co-founder/CEO, 2 Advance designees, 4 independent",
     "1 observer: Advance Publications (~30% owner). Farrell served 2021-2024, then became director.",
     "None disclosed in S-1"],
    ["Bumble\n(2021, PE)",
     "11 directors: founder/CEO, 2 Blackstone designees, 8 independent/affiliated. 73% women.",
     "1 observer: Blackstone (Martin Brand, Sr. MD). Elevated to director Aug 2024.",
     "None disclosed in S-1"],
    ["Rivian\n(2021)",
     "7 directors: founder/CEO, Amazon designee (Krawiec), 5 independent",
     "Observers: Amazon + Ford (strategic investors). Ford exited pre-IPO; Amazon retained post-IPO.",
     "None disclosed in S-1"],
    ["DoorDash\n(2020)",
     "9 directors: 3 co-founders, 3 VC directors (Kleiner, Sequoia, Silver Lake), 3 independent",
     "Per IRA: additional VC co-investors as observers. All terminated at IPO.",
     "None disclosed in S-1"],
]

for r_idx, row_data in enumerate(data):
    for c_idx, cell_text in enumerate(row_data):
        cell = table.rows[r_idx + 1].cells[c_idx]
        cell.text = cell_text
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.size = Pt(8)
                run.font.name = "Times New Roman"

# Move table before insert point
table_elem = table._element
body.remove(table_elem)
insert_point.addprevious(table_elem)

# Add interpretive paragraph
add_before(
    "Several patterns emerge from these examples. First, directors are consistently founders, "
    "lead VC partners, and experienced independent directors with operational expertise \u2014 "
    "people who have committed to fiduciary responsibility. Second, observers are typically "
    "additional VC investors (co-investors who did not receive a board seat), strategic "
    "investors (Amazon, Ford), or PE sponsors maintaining parallel information channels "
    "(Blackstone). Third, formal advisory boards are rare in S-1 filings \u2014 only Coinbase "
    "established one, and notably it was created post-IPO with regulatory and policy figures, "
    "not governance participants. Fourth, the observer-to-director pipeline is active: "
    "Andreessen (Coinbase), Farrell (Reddit), and Brand (Bumble) all transitioned from "
    "observer to director, suggesting that observer seats serve as a proving ground for "
    "eventual board membership. Fifth, observer rights almost universally terminate at IPO "
    "due to Reg FD concerns, with exceptions only for very large strategic shareholders "
    "(Amazon at Rivian, Advance at Reddit)."
)

# Save
output = "C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx"
doc.save(output)
print(f"Saved updated v3 to: {output}")
