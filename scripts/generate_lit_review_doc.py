"""Generate Word document from the prior literature summary."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.15


def heading(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.name = 'Times New Roman'
    return h


def bold_para(bold_text, normal_text=""):
    p = doc.add_paragraph()
    r = p.add_run(bold_text)
    r.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p


def bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p


def add_quote(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.right_indent = Inches(0.5)
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(10)
    return p


# Read the markdown file
with open("C:/Users/hjung/Documents/Claude/CorpAcct/Data/quality_reports/prior_literature_summary.md",
          "r", encoding="utf-8") as f:
    md = f.read()

# =====================================================================
# TITLE PAGE
# =====================================================================
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Prior Literature Review:\nBoard Observers in VC-Backed Firms")
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run("Research Project: The Three-Tier Board\nDirectors, Observers, and the Accountability Gap")
run.font.size = Pt(13)
run.italic = True

doc.add_paragraph()
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run("Harp Jung\nHarvard University\nMarch 2026")
run.font.size = Pt(12)

doc.add_page_break()

# =====================================================================
# SECTION 1: ACADEMIC PAPERS
# =====================================================================
heading("Section 1: Academic Papers", 1)

# Paper 1: Packin & Alon-Beck
heading("1.1 Packin & Alon-Beck (2025 U. Ill. L. Rev.)", 2)
bold_para("Citation: ", "Nizan Geslevich Packin & Anat Alon-Beck, Board Observers, 2025 U. ILL. L. REV. 1507-1568 (2025).")

bold_para("Main Finding: ", "First academic article to systematically explore board observers in the VC/startup ecosystem. Documents that 82% of VC entities use observers (100% of those with AUM >$500M). Observers represent 'control without formal control' \u2014 influence without fiduciary responsibility.")

bold_para("Methodology: ", "Legal analysis + NVCA Q4 2023 CFO Survey (46 VC/AVC entities). Case study of Microsoft-OpenAI observer arrangement. Comparative law analysis (US, UK, Australia, Korea).")

bold_para("Key Data:")
bullet("82% (38 of 46) surveyed entities use board observers")
bullet("100% of entities with AUM > $500M appoint board observers")
bullet("10 entities: >75% of portfolio companies have observers")
bullet("8 entities plan to increase number of observers")
bullet("All traditional VCs unanimously use observers")

bold_para("Legal Framework:")
bullet("Observers appointed via contractual instruments (IRAs, term sheets, side letters)")
bullet("NO voting rights, NO fiduciary duties, contractual information rights")
bullet("Obasi v. Tibet Pharmaceuticals (3d Cir. 2019): observers are NOT directors for Section 11 liability")
bullet("Three factors: no voting, tenure not shareholder-dependent, no fiduciary duty")

bold_para("Five Incentives for Observer Seats:")
bullet("Corporate governance \u2014 oversight without full board seat", "1. ")
bullet("Liability minimization \u2014 avoiding fiduciary duties", "2. ")
bullet("Antitrust compliance \u2014 avoiding Clayton Act Section 8", "3. ")
bullet("CFIUS regulation \u2014 foreign investors avoiding security reviews", "4. ")
bullet("ERISA compliance \u2014 VCOC management rights", "5. ")

add_quote("\"Board observers act as intermediaries, bridging active investors and corporate boards and aligning them with shared strategic goals.\"")

bold_para("Relevance: ", "THE foundational paper. Establishes prevalence, legal framework, and explicitly calls for empirical research on governance outcomes \u2014 which is exactly what our paper provides.")

doc.add_page_break()

# Paper 2: Ewens & Malenko
heading("1.2 Ewens & Malenko (NBER WP 27769, revised 2024)", 2)
bold_para("Citation: ", "Michael Ewens & Nadya Malenko, Board Dynamics over the Startup Life Cycle, NBER Working Paper No. 27769 (2020, revised May 2024).")

bold_para("Main Finding: ", "At formation, boards are entrepreneur-controlled. Independent directors join after the second financing with a tie-breaking vote. They play a 'mediation' role (not just monitoring) between VCs and entrepreneurs. Control shifts to VCs at later stages.")

bold_para("Methodology: ", "Novel panel dataset: 7,780 startups over 2002-2017, built from SEC Form D filings + VentureSource. DiD design using AWS introduction (2006) as shock to VC capital requirements.")

bold_para("Key Data:")
bullet("Average board: 4.5 directors (2 VCs, 1.7 executives, 0.8 independent)")
bullet("First financing: 47% entrepreneur-controlled")
bullet("Fourth financing: 62% VC-controlled")
bullet("Independent directors present in ~50% of firm-year observations")

add_quote("\"The firm could establish an advisory board (as is frequently done in private firms) or add these individuals to the board as 'board observers.'\" (Footnote 3)")

bold_para("Relevance: ", "Essential empirical foundation. Footnote 3 explicitly distinguishes observers from directors. Their Form D methodology is directly applicable. Our CIQ observer data fills the gap they acknowledge.")

doc.add_page_break()

# Paper 3: Nili
heading("1.3 Nili (2025) \u2014 Interlocking Directorates", 2)
bold_para("Citation: ", "Yaron Nili, Interlocking Directorates in the United States, Ch. 15 in RESEARCH HANDBOOK ON COMPETITION AND CORPORATE LAW (2025).")

bold_para("Main Finding: ", "Despite Clayton Act Section 8 prohibitions, horizontal directorships are prevalent and increasing: 1,888 directors served on boards of multiple companies in the same industry in 2019.")

bold_para("Relevance: ", "Explains why observers are the Clayton Act workaround. Investors use observer seats instead of director seats to maintain presence on multiple boards without triggering Section 8. Directly motivates our Test 3 (information spillover) and the January 2025 DOJ/FTC enforcement action extending Section 8 to observers.")

doc.add_page_break()

# Paper 4: Carter
heading("1.4 Carter (JAE, 2025) \u2014 THE Key Citation", 2)
bold_para("Citation: ", "Mary Ellen Carter, What Director Experience Matters: A Discussion of Donelson, Hutzler and Rhodes (2025), J. ACCT. & ECON. 80, 101816 (2025).")

bold_para("Main Finding: ", "Discussion paper in a top accounting journal that EXPLICITLY calls for future research on board observers.")

add_quote("\"Board observers may provide a conduit for information, yielding benefits similar to competitor interlocking directors and possibly remain under the DOJ radar. Future research might examine how prevalent they are in public companies and how their presence impacts corporate governance and firm performance.\"")

bold_para("Relevance: ", "This is the direct academic mandate for our paper. Carter identifies exactly the research question we answer: (1) observer prevalence (our CIQ/EDGAR data), (2) information conduit function (our Test 3 spillover results, p<0.001), (3) governance impact (our Test 5 results, p<0.001 for lawsuits). Our finding that non-fiduciary connections transmit 3x more information (p=0.016) directly confirms Carter's hypothesis about observers as information conduits.")

doc.add_page_break()

# Paper 5: Pollman
heading("1.5 Pollman (2024) \u2014 Theoretical Framework", 2)
bold_para("Citation: ", "Elizabeth Pollman, Dynamic Views of Startup Governance and Failure, U. Penn. ILE Research Paper No. 24-15 (2024).")

bold_para("Main Finding: ", "Synthesizes three perspectives on startup governance: (1) monitoring model, (2) horizontal/dynamic governance, (3) power law portfolio view. Governance involves heterogeneous shareholders in overlapping roles with vertical AND horizontal tensions.")

bold_para("Relevance: ", "Provides the theoretical scaffolding for why the three-tier architecture exists. As startups add financing rounds, governance complexity increases. Observers are a natural response \u2014 including additional voices without expanding the formal board.")

doc.add_page_break()

# =====================================================================
# SECTION 2: PRACTITIONER SOURCES
# =====================================================================
heading("Section 2: Practitioner Sources (Mark Suster Blog Series)", 1)

doc.add_paragraph("Mark Suster, Managing Partner of Upfront Ventures, published a series of blog posts on startup boards in February 2019. These provide the most detailed practitioner perspective on board observers available.")

heading("Practitioner Taxonomy of Board Observers", 2)
doc.add_paragraph("Suster identifies four types of board observers:")

bold_para("Type 1: VC 'Plus-1s'. ", "Partners bring analysts or associates to observe board meetings. Used for training junior staff and ensuring follow-up on action items. The most common and least controversial type.")

bold_para("Type 2: Strategic Investors. ", "Corporate VCs (CVCs) who prefer not to take formal board seats (often due to corporate policy) or who receive observer seats because the startup wants to limit their influence.")

bold_para("Type 3: Minority Shareholders. ", "Late-round investors with smaller check sizes who request observer seats as a compromise \u2014 they want board-level access but don't warrant a full seat given their investment size.")

bold_para("Type 4: VCs Avoiding Board Obligations. ", "Investors who want board influence without legal responsibilities and liabilities. Suster views this as the most problematic type: 'Either they're committed and they should take a board seat with full rights and responsibilities, or they should trust others who are willing to do the actual board work.'")

heading("Key Practitioner Insights", 2)

bold_para("On observer influence: ", "'Almost nothing controversial is EVER voted on at a startup board meeting. Most votes are for administrative tasks. So the value of a person in the room is the value of their speaking and ability either to persuade others or to be disruptive when you're seeking consensus.'")

bold_para("On practical equivalence: ", "'For all intents and purposes a board observer is like a full board member so tread lightly in handing out these seats.'")

bold_para("On board degradation: ", "'Can I really manage a 7-person board plus 2 board observers?' Observers inflate the effective board size, creating scheduling and discussion management challenges.")

bold_para("On low-functioning boards: ", "'Sometimes boards have members with no skin in the game and no personal responsibility in outcomes.' Observers fit this description exactly \u2014 full information access, no fiduciary duty, no personal liability.")

bold_para("Recommendation: ", "If observers are unavoidable, insist on 'silent observer' agreements. Observers should truly observe \u2014 not participate unless called upon.")

doc.add_page_break()

# =====================================================================
# SECTION 3: MAPPING TO RESEARCH
# =====================================================================
heading("Section 3: How This Literature Maps to Our Research", 1)

heading("3.1 Empirical Gaps Identified", 2)

doc.add_paragraph("The literature identifies six empirical gaps that our paper addresses:")

bullet("Observer prevalence at scale \u2014 Packin & Alon-Beck provide survey data (82%), but no large-sample analysis. Our CIQ data: 3,058 companies, 4,915 observers.", "Gap 1: ")
bullet("Impact on governance outcomes \u2014 No existing study. Our Test 5: lawsuits OR=0.46 (p<0.001), acquisitions OR=0.51 (p=0.002).", "Gap 2: ")
bullet("Information flow through observer networks \u2014 Carter (2025) hypothesizes; no evidence. Our Test 3: same-industry spillover +0.58% (p<0.001).", "Gap 3: ")
bullet("Accountability structure effects \u2014 Does fiduciary status matter? Our Test 3: non-fiduciary connections transmit 3x more information (p=0.016).", "Gap 4: ")
bullet("Observer vs. director tradeoff \u2014 What determines role assignment? Our S-1 exhibit data codes observer provisions in 320 IRAs.", "Gap 5: ")
bullet("Chilling effect and informal influence \u2014 Packin & Alon-Beck theorize; Suster describes operationally. No empirical evidence yet.", "Gap 6: ")

heading("3.2 Theoretical Lenses", 2)

bold_para("1. Monitoring/Advising vs. Mediation (Ewens & Malenko): ", "Independent directors mediate VC-entrepreneur conflicts via tie-breaking votes. Observers serve a different function \u2014 information channels and informal advisors without formal mediation power.")

bold_para("2. Control Without Formal Control (Packin & Alon-Beck): ", "Observers allow investors to maintain influence without triggering fiduciary duties, antitrust restrictions, or CFIUS reviews. Real governance effects despite lacking voting power.")

bold_para("3. Dynamic Governance Complexity (Pollman): ", "As financing rounds multiply, governance complexity increases. Observers are a natural mechanism to include additional voices without expanding the formal board.")

heading("3.3 Key Academic Mandate", 2)

doc.add_paragraph("Carter (JAE, 2025) provides the most direct academic mandate:")

add_quote("\"Board observers may provide a conduit for information, yielding benefits similar to competitor interlocking directors and possibly remain under the DOJ radar. Future research might examine how prevalent they are in public companies and how their presence impacts corporate governance and firm performance.\"")

doc.add_paragraph("Our paper answers every part of this call: prevalence (CIQ + S-1 data), information conduit (Test 3 spillover results), governance impact (Tests 1 and 5), and DOJ implications (Clayton Act 2025 analysis).")

# =====================================================================
# SAVE
# =====================================================================
output_path = "C:/Users/hjung/Documents/Claude/CorpAcct/Prior Literature Review -- Board Observers.docx"
doc.save(output_path)
print(f"Saved to: {output_path}")
