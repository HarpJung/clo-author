"""Generate Word document: Four Empirical Settings for Three-Tier Board Governance Research."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

doc = Document()

# --- Styles ---
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

# Helper functions
def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.style = doc.styles['Normal']
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r + 1].cells[c]
            cell.text = val
            for p in cell.paragraphs:
                p.style = doc.styles['Normal']
                for run in p.runs:
                    run.font.size = Pt(10)
    return table

def add_bold_para(bold_text, normal_text=""):
    p = doc.add_paragraph()
    run_b = p.add_run(bold_text)
    run_b.bold = True
    if normal_text:
        p.add_run(normal_text)
    return p

def add_bullet(text, bold_prefix=""):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run_b = p.add_run(bold_prefix)
        run_b.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

# =====================================================================
# TITLE PAGE
# =====================================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Four Empirical Settings for Studying\nthe Three-Tier Board Governance Architecture")
run.bold = True
run.font.size = Pt(18)
run.font.name = 'Times New Roman'

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Directors, Observers, and Advisors:\nFrom Qualitative Evidence to Causal Identification")
run.font.size = Pt(14)
run.font.name = 'Times New Roman'
run.italic = True

doc.add_paragraph()
author = doc.add_paragraph()
author.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = author.add_run("Harp Jung\nHarvard University")
run.font.size = Pt(12)

date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_para.add_run("March 2026")
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()
note = doc.add_paragraph()
note.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = note.add_run("Research Proposal — Future Empirical Agenda")
run.italic = True
run.font.size = Pt(11)

doc.add_page_break()

# =====================================================================
# EXECUTIVE SUMMARY
# =====================================================================
add_heading_styled("Executive Summary", level=1)

doc.add_paragraph(
    "This document outlines four empirical research settings for studying the three-tier "
    "board governance architecture — the system of directors, observers, and advisors that "
    "governs most venture-backed private firms. Each setting offers a distinct identification "
    "strategy for testing how this governance configuration affects corporate accountability."
)

doc.add_paragraph(
    "The settings are ordered by data feasibility, from fully public data to proprietary partnerships:"
)

add_bullet("S-1 Mining — public data, hand-collected from SEC EDGAR", "Setting 1: ")
add_bullet("NVCA 2020 Model Document Change — natural experiment, partially public", "Setting 2: ")
add_bullet("DOJ/FTC 2025 Clayton Act Statement — regulatory shock, partially public", "Setting 3: ")
add_bullet("D&O Insurance and the Accountability Vacuum — requires proprietary data partnerships", "Setting 4: ")

doc.add_paragraph(
    "For the class write-up, we frame the descriptive three-tier argument with interview evidence "
    "and propose all four settings as the future research agenda. For journal publication, Setting 1 "
    "(S-1 mining) requires only public data and should be pursued first."
)

doc.add_page_break()

# =====================================================================
# SETTING 1: S-1 MINING
# =====================================================================
add_heading_styled("Setting 1: S-1 Mining — Board Governance Architecture and Post-IPO Outcomes", level=1)

add_heading_styled("The Mechanism", level=2)

doc.add_paragraph(
    "When a VC-backed company goes public, it files an S-1 registration statement with the SEC. "
    "The S-1 is extraordinarily detailed about governance arrangements — including board composition, "
    "investor rights agreements, observer provisions, advisory boards, and D&O coverage. Critically, "
    "the S-1 captures the pre-IPO governance architecture — the three-tier system as it existed "
    "during the company's private life."
)

doc.add_paragraph("The mechanism is:")
add_bullet("Private firms develop governance architectures (directors + observers + advisors) during VC funding rounds")
add_bullet("These architectures create varying levels of accountability (the core argument: the tiers differ in personal risk)")
add_bullet("The governance \"habits\" and reporting quality established under this architecture persist into the post-IPO period")
add_bullet(
    "Firms with weaker pre-IPO accountability structures (more observer reliance, less director "
    "oversight, D&O-insulated directors) should exhibit worse post-IPO governance outcomes"
)

doc.add_paragraph(
    "Why this works: The S-1 creates a moment of mandatory, detailed disclosure about a private "
    "firm's governance — a window into an otherwise opaque system. You observe the governance "
    "architecture at IPO and then track outcomes."
)

add_heading_styled("Specific Hypotheses", level=2)

add_bold_para("H1 (Observer reliance): ",
    "Firms with a higher ratio of pre-IPO observer seats to director seats exhibit worse post-IPO "
    "financial reporting quality (more restatements, higher accruals, higher audit fees).")
add_bullet(
    "Mechanism: If observers have full information but no duty to act, they won't intervene even "
    "when they see reporting problems. The board may \"know\" about issues but lack a participant "
    "with both knowledge AND obligation to flag them."
)

add_bold_para("H2 (Board degradation): ",
    "Firms that went through more funding rounds (and thus accumulated more observers) have lower "
    "CEO turnover-performance sensitivity post-IPO.")
add_bullet(
    "Mechanism: Matt's observation — as the board grows from 3-4 to 12-15, governance shifts from "
    "discussion to presentation. This inertia persists."
)

add_bold_para("H3 (D&O and monitoring): ",
    "Firms where directors were fully D&O-insulated pre-IPO (disclosed in S-1) exhibit weaker "
    "monitoring outcomes — consistent with moral hazard at the director tier completing the "
    "accountability vacuum.")

add_bold_para("H4 (Advisory board effect): ",
    "Firms with formal advisory boards show different governance outcomes — either better (access "
    "to expertise) or worse (diffusion of accountability). Directional prediction depends on theory.")

add_heading_styled("Data Required", level=2)

add_table(
    ["Data Element", "Source", "Collection Method"],
    [
        ["Observer seats, advisory boards, board composition pre-IPO",
         "S-1 filings on SEC EDGAR",
         "Hand-collect or NLP/text search for \"observer,\" \"advisory board,\" \"board of advisors,\" \"non-voting\" in S-1 risk factor and governance sections"],
        ["D&O insurance details",
         "S-1 filings (often in \"Executive Compensation\" or \"Description of Capital Stock\" sections)",
         "Hand-collect; look for indemnification agreements, D&O coverage amounts"],
        ["Investor rights agreements",
         "S-1 exhibits (Investor Rights Agreement is typically filed as an exhibit)",
         "Hand-collect observer rights clauses from exhibits"],
        ["Post-IPO restatements",
         "Audit Analytics Restatement database",
         "Merge by CIK/ticker"],
        ["Post-IPO fraud",
         "SEC enforcement actions (AAER database), Stanford Securities Class Action Clearinghouse",
         "Standard merge"],
        ["Post-IPO performance",
         "CRSP (returns), Compustat (accounting)",
         "Standard merge"],
        ["CEO turnover",
         "ExecuComp or hand-collected from proxy statements",
         "Standard merge"],
        ["Audit fees / audit quality",
         "Audit Analytics",
         "Standard merge"],
        ["VC characteristics",
         "VentureSource, PitchBook, or Ewens & Malenko GitHub data",
         "Merge by company name/CIK"],
    ]
)

add_heading_styled("Sample Construction", level=2)

add_bullet("Universe: All VC-backed IPOs on U.S. exchanges, 2004-2024 (approximately 1,500-2,000 firms)")
add_bullet("S-1s available via EDGAR full-text search")
add_bullet("Programmatically search for \"observer\" OR \"advisory board\" OR \"non-voting\" in S-1 filings")
add_bullet("Hand-code the governance architecture for the subset that mentions these terms")
add_bullet("Control group: VC-backed IPOs with NO mention of observer/advisor arrangements")

add_heading_styled("Feasibility Assessment", level=2)

add_bold_para("Strengths: ",
    "All data is public. S-1 is the richest governance disclosure document for private firms. "
    "Post-IPO outcome data is well-established. No proprietary data needed.")

add_bold_para("Challenges: ",
    "Hand-collection is labor-intensive (though NLP can pre-screen). Observer/advisor disclosure "
    "is not standardized — some firms disclose more than others. Selection: firms that disclose "
    "observer arrangements may differ from those that don't in unobservable ways.")

add_bold_para("Identification: ",
    "Pre-IPO governance is determined years before the IPO (at each funding round), providing "
    "temporal separation from post-IPO outcomes. Could instrument with VC firm fixed effects "
    "(some VCs always take observer seats, others always take board seats — Matt describes "
    "this as a firm-level policy choice).")

doc.add_page_break()

# =====================================================================
# SETTING 2: NVCA 2020
# =====================================================================
add_heading_styled("Setting 2: NVCA 2020 Model Document Change (Natural Experiment)", level=1)

add_heading_styled("The Mechanism", level=2)

doc.add_paragraph(
    "In 2020, the NVCA revised its model legal documents — the standard templates used in "
    "approximately 85% of VC financing transactions by 2022 (Bartlett et al. 2023, Journal of "
    "Law and Economics, using ~5,000 Delaware startup charters). The key change: removal of "
    "language requiring board observers to act \"in a fiduciary manner.\""
)

add_bold_para("Before 2020: ",
    "The NVCA Investor Rights Agreement stated observers must act \"in a fiduciary manner\" "
    "regarding confidential information. This language was already anomalous — observers have "
    "no fiduciary duties under Delaware law — but it created at least a contractual expectation "
    "of fiduciary-like behavior.")

add_bold_para("After 2020: ",
    "This language was deleted. Confidentiality obligations for observers now depend entirely on "
    "separate contractual provisions (NDAs, side letters). The removal was described by "
    "practitioners as a \"correction\" — aligning the model documents with the legal reality "
    "that observers owe no fiduciary duties.")

doc.add_paragraph("The mechanism:")
add_bullet("Pre-2020, the fiduciary language created a behavioral anchor — observers who signed agreements containing \"fiduciary manner\" language may have behaved more cautiously with confidential information")
add_bullet("Post-2020, removing this anchor formally liberated observers from any fiduciary expectation")
add_bullet("If observers' behavior changes in response (even at the margin), governance outcomes should shift")

add_bold_para("Key paper to cite: ",
    "Bartlett, Gulati, and Liskow (2023), Journal of Law and Economics, 55(1). They analyzed "
    "~5,000 charters and found NVCA adoption went from 3% (2004) to 85% (2022). Their methodology — "
    "cosine similarity analysis of charter text against the NVCA model — provides a template for "
    "identifying which firms adopted the 2020 changes.")

add_heading_styled("Specific Hypotheses", level=2)

add_bold_para("H1 (Information leakage): ",
    "After the 2020 removal, competitively sensitive information is more likely to flow between "
    "portfolio companies sharing common observers — proxied by changes in competitor behavior, "
    "pricing, or product launches.")

add_bold_para("H2 (Governance quality): ",
    "Firms funded post-2020 under the new NVCA terms exhibit weaker governance outcomes (more "
    "fraud, lower reporting quality) because the last nominal accountability constraint on "
    "observers was removed.")

add_bold_para("H3 (D&O demand): ",
    "Post-2020, demand for D&O insurance among observers increases — the paradox Packin & "
    "Alon-Beck note. Removing fiduciary language may make observers more anxious about liability "
    "(because the contractual ambiguity is resolved against them), increasing D&O demand.")

add_heading_styled("Data Required", level=2)

add_table(
    ["Data Element", "Source", "Notes"],
    [
        ["Charter text / Investor Rights Agreement text",
         "Delaware Division of Corporations (charters); S-1 exhibits (for IPO firms)",
         "Bartlett et al. (2023) used ~5,000 charters from Delaware. Their methodology is replicable."],
        ["NVCA adoption indicator",
         "Apply Bartlett et al.'s cosine similarity methodology to post-2020 charters",
         "Identify which firms adopted the revised NVCA language"],
        ["Observer-specific provisions",
         "Investor Rights Agreements (filed as S-1 exhibits for IPO firms; otherwise proprietary)",
         "Pre-2020: search for \"fiduciary manner\"; Post-2020: search for absence"],
        ["Governance outcomes",
         "Same as Setting 1 for public firms; for private firms, need VentureSource/PitchBook exit data",
         ""],
        ["Fraud/distress",
         "SEC enforcement (public); Dyck et al. (2026) methodology for private firm fraud",
         ""],
    ]
)

add_heading_styled("Key Data Challenge", level=2)

doc.add_paragraph(
    "The Investor Rights Agreement (which contains the observer provisions) is only publicly "
    "available when filed as an S-1 exhibit at IPO. For firms that stay private, this document "
    "is confidential. Options:")
add_bullet("Restrict to IPO firms and use the S-1 exhibit (public data, same as Setting 1)", "1. ")
add_bullet("Partner with law firms that draft these agreements — Bartlett et al. worked with Cooley, Wilson Sonsini, and other major VC law firms to access charter data", "2. ")
add_bullet("Use the NVCA directly — Packin & Alon-Beck partnered with NVCA for their survey data; a follow-up study could be proposed", "3. ")

add_heading_styled("Feasibility Assessment", level=2)

add_bold_para("Strengths: ",
    "Clean treatment (specific language change on a specific date). High adoption rate (85%) "
    "means treatment is pervasive. Bartlett et al. provide the methodological template.")

add_bold_para("Challenges: ",
    "The key variable (observer fiduciary language) is in a document that's only public at IPO. "
    "For a DiD, you need pre-2020 and post-2020 observations — firms that IPO'd after 2020 with "
    "post-2020 NVCA terms vs. firms that IPO'd before 2020 with pre-2020 terms. Potential "
    "confounds: COVID, market conditions, other NVCA changes in 2020.")

doc.add_page_break()

# =====================================================================
# SETTING 3: DOJ/FTC 2025
# =====================================================================
add_heading_styled("Setting 3: DOJ/FTC 2025 Clayton Act Statement (Regulatory Shock)", level=1)

add_heading_styled("The Mechanism", level=2)

doc.add_paragraph(
    "On January 10, 2025, the DOJ and FTC filed a joint Statement of Interest in Musk v. Altman "
    "(N.D. Cal.) arguing that board observers are subject to Section 8 of the Clayton Act's "
    "prohibition on interlocking directorates — the same rules that apply to formal directors. "
    "This was adopted unanimously, including by Trump's FTC designee."
)

add_bold_para("Before January 2025: ",
    "Board observers were explicitly not covered by Section 8. The FTC itself had stated: "
    "\"Board observers are not subject to the Section 8 ban on interlocking directorates.\" "
    "Investors used this as a loophole — placing observers on competitor boards to gain "
    "information without triggering antitrust scrutiny.")

add_bold_para("After January 2025: ",
    "The agencies argued that observers who attend meetings at competing companies have \"the "
    "same anticompetitive access as an interlocking director.\" This dramatically increased the "
    "regulatory cost of observer arrangements, particularly for: VC/PE firms with portfolio "
    "companies in overlapping markets; CVC investors observing boards of potential competitors; "
    "any investor with observer seats at firms with competitive overlap.")

doc.add_paragraph("The mechanism:")
add_bullet("Pre-2025, observer seats at competitors were low-cost (no Section 8 liability)")
add_bullet("The DOJ/FTC statement created a sudden increase in the regulatory cost of maintaining observer seats at competing firms")
add_bullet("Firms with observer-based interlocks should respond: observer resignations, board restructuring, changes in VC investment patterns")
add_bullet("If observers were providing governance value (monitoring, advising), their removal should worsen governance outcomes — the trade-off between antitrust risk and governance quality")

add_heading_styled("Critical Prior Papers", level=2)

add_bold_para("Donelson, Hutzler & Rhodes (2025), ",
    "\"Does Antitrust Enforcement Against Interlocking Directorates Impair Corporate Governance?,\" "
    "Journal of Accounting and Economics. This paper ALREADY did an event study around the October "
    "2022 DOJ enforcement wave (director resignations from Thoma Bravo, Prosus, etc.). They found:")
add_bullet("Competitor-interlocked directors were more likely to leave boards")
add_bullet("Replacements had less industry experience (47 years average departing vs. much less for replacements)")
add_bullet("This weakened governance: interlocked directors with industry experience produced higher profit margins and were more likely to fire underperforming CEOs")
add_bullet("Sample: 1.6 million director-company-month observations, 2004-2022")

doc.add_paragraph(
    "The extension: Donelson et al. studied the October 2022 shock to director interlocks. "
    "The January 2025 shock is specific to observer interlocks — a different treatment affecting "
    "a different governance tier. Their methodology can be replicated for the observer-specific shock."
)

add_bold_para("Poberejsky (2025, JMP), ",
    "\"Interlocking Directorates, Competition, and Innovation,\" Northwestern/Cornerstone Research "
    "(SSRN: 4944799). Found horizontal interlocks reduce competition but increase innovation "
    "(17% quantity, 30% quality). Mechanism: market segmentation. Uses public firm data; does not "
    "cover observers.")

add_bold_para("Nili et al. (2025), ",
    "published in research handbook and presented at Harvard Law Forum. Found 2,309 individual "
    "interlocks and 2,927 investor-level interlocks. 65% of interlocking directors hold positions "
    "at PE/VC firms investing in the competing companies. Covers both public and private firms — "
    "this is the dataset you'd want access to.")

add_heading_styled("Specific Hypotheses", level=2)

add_bold_para("H1 (Observer departures): ",
    "After January 2025, observer seats at firms with competitive overlap decline — observers "
    "resign or firms restructure to eliminate the interlock.")

add_bold_para("H2 (Governance trade-off): ",
    "Firms that lose observers due to the regulatory shock experience worse governance outcomes — "
    "weakened monitoring, less information flow, slower intervention in distress. "
    "(Parallels Donelson et al.'s finding for directors.)")

add_bold_para("H3 (Substitution to advisors): ",
    "Investors who lose observer seats at competitors may substitute to advisory board roles — "
    "which are even further from Section 8's reach. The three-tier system adapts: regulatory "
    "pressure on Tier 2 (observers) pushes activity to Tier 3 (advisors).")

add_bold_para("H4 (VC portfolio reshuffling): ",
    "VC/PE firms with observer-based interlocks restructure their portfolios — divesting from "
    "companies where the interlock creates antitrust exposure.")

add_heading_styled("Data Required", level=2)

add_table(
    ["Data Element", "Source", "Notes"],
    [
        ["Observer seats at competing firms",
         "PitchBook (board composition + industry classification)",
         "PitchBook tracks some board observers; need to verify coverage"],
        ["Competitive overlap identification",
         "SIC/NAICS codes, patent overlap, product market similarity",
         "Donelson et al. use TNIC (text-based network industry classification by Hoberg & Phillips)"],
        ["Observer departures/appointments",
         "PitchBook board changes; SEC 8-K filings",
         "For public firms: 8-K filings. For private: PitchBook or hand-collection."],
        ["VC portfolio data",
         "PitchBook, VentureSource, Crunchbase",
         "Identify which VC firms have observers at competing portfolio companies"],
        ["Governance outcomes",
         "Same as Settings 1 & 2",
         ""],
        ["Event date",
         "January 10, 2025 (DOJ/FTC Statement of Interest filing)",
         "Clean event"],
        ["Nili et al. interlock dataset",
         "Contact authors (Duke/BU law)",
         "2,309 individual + 2,927 investor-level interlocks; covers public AND private firms"],
    ]
)

add_heading_styled("Feasibility Assessment", level=2)

add_bold_para("Strengths: ",
    "Clean exogenous shock (government action, not firm choice). Prior paper (Donelson et al.) "
    "provides exact methodology to replicate. Extends to observers (a genuinely new treatment group). "
    "Nili et al. have the most comprehensive interlock dataset including private firms.")

add_bold_para("Challenges: ",
    "Observer data for private firms is not systematically captured in any database. PitchBook "
    "has some observer data but coverage is uncertain. The January 2025 statement is recent, so "
    "post-shock outcomes are still short-horizon. Public firm observers are rare (Packin & "
    "Alon-Beck note observers are \"exceedingly rare\" in public companies due to Reg FD).")

add_bold_para("Best approach: ",
    "Focus on late-stage private firms approaching IPO (where competitive overlap is identifiable "
    "and observer arrangements are more likely to be documented in subsequent S-1 filings).")

doc.add_page_break()

# =====================================================================
# SETTING 4: D&O INSURANCE
# =====================================================================
add_heading_styled("Setting 4: D&O Insurance and the Accountability Vacuum", level=1)

add_heading_styled("The Mechanism", level=2)

doc.add_paragraph("The interviews reveal that D&O insurance completes the accountability vacuum at the director tier:")
add_bullet("Toj: D&O is \"almost a prerequisite to being on a board\"")
add_bullet("Matt: pays ~$37K/year for personal D&O coverage; fund also has coverage")
add_bullet("The interviewer asks: \"Does [D&O] make you care less?\" Matt: \"Maybe... yeah, that's actually a hard way to put it, but... it allows you to care less.\"")

doc.add_paragraph("The mechanism:")
add_bullet("Tier 3 (Advisors): No fiduciary duties \u2192 no accountability", "1. ")
add_bullet("Tier 2 (Observers): No fiduciary duties \u2192 no accountability (confirmed by NVCA 2020)", "2. ")
add_bullet("Tier 1 (Directors): Have fiduciary duties, BUT D&O insurance transfers personal financial risk to insurers \u2192 accountability is nominally present but economically neutralized", "3. ")
add_bullet("Result: No participant in the governance architecture bears meaningful personal financial risk from governance failures", "4. ")

doc.add_paragraph(
    "This is the most provocative version of the argument. The testable question is: Does D&O "
    "coverage reduce director monitoring effort? If yes, the accountability gap runs all the way "
    "through — not just tiers 2 and 3."
)

add_heading_styled("The Data Problem (and Solutions)", level=2)

doc.add_paragraph(
    "The core problem: D&O insurance data for private firms is not publicly available. All "
    "existing empirical D&O research uses either Canadian public firms (where D&O disclosure "
    "is mandatory) or regulatory shocks to D&O policy features."
)

add_table(
    ["Data Source", "What It Contains", "Access", "Papers That Use It"],
    [
        ["Canadian D&O disclosures",
         "D&O premiums, limits, deductibles, coverage details for public firms",
         "Public (SEDAR filings)",
         "Core (2000), Chalmers et al. (2002), Lin et al. (2011, 2012, 2019), Boyer & Stern (2012)"],
        ["Advisen Loss Data",
         "Proprietary database of D&O claims and losses, 93 countries, 20+ years",
         "Proprietary",
         "Used by insurance industry; some academic papers reference it"],
        ["Aon/Marsh/WTW broker data",
         "D&O premium benchmarks, coverage trends, claim frequency",
         "Proprietary — available through broker relationships",
         "Baker & Griffith (2010), Ensuring Corporate Misconduct"],
        ["D&O Diary / Woodruff Sawyer data",
         "Public commentary and data on D&O pricing trends, claim activity",
         "Semi-public (blog + reports)",
         "Kevin LaCroix's D&O Diary publishes annual claim data summaries"],
        ["SEC S-1 filings",
         "Some S-1s disclose D&O coverage amounts and indemnification provisions",
         "Public",
         "Not systematically studied for private-to-public firms"],
    ]
)

add_heading_styled("Key Prior Papers Using Proprietary D&O Data", level=2)

add_bold_para("Baker & Griffith (2010), ",
    "Ensuring Corporate Misconduct: How Liability Insurance Undermines Shareholder Litigation, "
    "University of Chicago Press. Book-length study using proprietary D&O underwriting data from "
    "a major insurance broker. Gold standard for understanding D&O in practice but data is proprietary.")

add_bold_para("Gillan & Panasian (2015), ",
    "\"On Lawsuits, Corporate Governance, and the Dodd-Frank Act,\" Journal of Corporate Finance. "
    "Uses a combination of Advisen loss data and NAIC data to study D&O litigation risk.")

add_heading_styled("Feasible Research Designs", level=2)

add_bold_para("Option A: Extend the Canadian D&O literature to the three-tier question. ",
    "Canada requires D&O disclosure for public firms. Use Canadian D&O data to study whether "
    "higher D&O coverage reduces CEO turnover-performance sensitivity (the accountability mechanism). "
    "This extends Lin et al. (2011) with the theoretical framing — D&O is one of three mechanisms "
    "that neutralize accountability, alongside observer status (no duty) and advisor status (no role).")

add_bold_para("Option B: Use S-1-disclosed D&O data for VC-backed IPO firms. ",
    "Many S-1 filings disclose D&O insurance details. Combine with Setting 1: for each VC-backed "
    "IPO, code (i) the governance architecture AND (ii) the D&O coverage. Test whether the "
    "combination of strong D&O + heavy observer reliance predicts worse post-IPO outcomes than "
    "either alone. This is the interaction effect: D&O alone doesn't destroy accountability; "
    "observers alone don't destroy accountability; but the combination creates the full vacuum.")

add_bold_para("Option C: Partner with a D&O broker. ",
    "Several academic papers have obtained proprietary D&O data through broker partnerships. "
    "Woodruff Sawyer specifically publishes on VC/startup D&O. Priya Cherian Huskins at Woodruff "
    "Sawyer has written about whether board observers should get D&O coverage — she could be a "
    "data partner.")

add_heading_styled("Specific Hypotheses", level=2)

add_bold_para("H1 (D&O moral hazard): ",
    "Higher D&O coverage is associated with lower director monitoring intensity (fewer executive "
    "dismissals, higher discretionary accruals, lower earnings quality). Extension of Lin et al. (2011).")

add_bold_para("H2 (The accountability interaction): ",
    "The negative governance effect of D&O is amplified when combined with a heavy observer "
    "presence — because D&O removes the last remaining accountability mechanism. "
    "Test the interaction: D&O coverage \u00d7 Observer ratio \u2192 Governance outcome.")

add_bold_para("H3 (D&O and role selection): ",
    "In settings where D&O is unavailable or expensive, investors are more likely to take observer "
    "seats rather than director seats (to avoid uninsured fiduciary exposure). D&O availability "
    "causes more director seats. This would be the first paper to study D&O as a determinant of "
    "governance architecture, not just a consequence of it.")

add_heading_styled("Feasibility Assessment", level=2)

add_bold_para("Strengths: ",
    "Theoretically the most novel setting — connects D&O (well-studied for public firms) to the "
    "three-tier architecture (unstudied). The interaction hypothesis (D&O \u00d7 observer ratio) "
    "is genuinely new.")

add_bold_para("Challenges: ",
    "Private firm D&O data is the binding constraint. Option B (S-1 data) is most feasible but "
    "limits you to IPO firms. Option C (broker partnership) is highest payoff but requires "
    "relationship-building.")

doc.add_page_break()

# =====================================================================
# SUMMARY TABLE
# =====================================================================
add_heading_styled("Summary: Data Requirements Across All Four Settings", level=1)

add_table(
    ["Setting", "Public Data Available?", "Proprietary Data Needed?", "Who Has the Proprietary Data?"],
    [
        ["1. S-1 Mining",
         "Yes — S-1 on EDGAR, outcome data from standard sources",
         "No — fully public",
         "N/A"],
        ["2. NVCA 2020 Change",
         "Partially — S-1 exhibits for IPO firms",
         "Yes — Investor Rights Agreements for private firms",
         "VC law firms (Cooley, Wilson Sonsini, Gunderson); NVCA; Bartlett et al. may share charter data"],
        ["3. DOJ/FTC 2025 Shock",
         "Partially — public firm interlocks identifiable",
         "Yes — private firm observer data",
         "PitchBook (some coverage); Nili et al. interlock dataset (Duke/BU); NVCA"],
        ["4. D&O Accountability",
         "Partially — Canadian public; S-1 disclosures",
         "Yes — private firm D&O premiums and terms",
         "D&O brokers (Woodruff Sawyer, Marsh, Aon); Advisen loss database; NAIC"],
    ]
)

doc.add_paragraph()

# =====================================================================
# RECOMMENDED SEQUENCING
# =====================================================================
add_heading_styled("Recommended Sequencing", level=1)

add_bold_para("For the class project: ",
    "Frame the descriptive three-tier argument with interview evidence and propose all four "
    "settings as the future research agenda.")

add_bold_para("For the journal paper: ",
    "Start with Setting 1 (S-1 mining) — it requires only public data, creates a novel "
    "hand-collected dataset, and directly tests the three-tier architecture's consequences. "
    "This alone is a publishable paper. Settings 2-4 become follow-up papers as you build "
    "data relationships.")

doc.add_paragraph()

# =====================================================================
# KEY REFERENCES
# =====================================================================
add_heading_styled("Key References", level=1)

refs = [
    "Adams, R. B., & Ferreira, D. (2007). A theory of friendly boards. Journal of Finance, 62(1), 217-250.",
    "Adams, R. B., Hermalin, B. E., & Weisbach, M. S. (2010). The role of boards of directors in corporate governance: A conceptual framework and survey. Journal of Economic Literature, 48(1), 58-107.",
    "Amornsiripanitch, N., Gompers, P. A., & Xuan, Y. (2019). More than money: Venture capitalists on boards. Journal of Law, Economics, and Organization, 35(3), 513-555.",
    "Baker, T., & Griffith, S. J. (2010). Ensuring Corporate Misconduct: How Liability Insurance Undermines Shareholder Litigation. University of Chicago Press.",
    "Bartlett, R. P., Gulati, M., & Liskow, E. (2023). Standardization and innovation in venture capital contracting. Journal of Law and Economics, 55(1).",
    "Bebchuk, L. A., & Fried, J. M. (2004). Pay Without Performance: The Unfulfilled Promise of Executive Compensation. Harvard University Press.",
    "Boyer, M. M., & Stern, L. H. (2012). Is corporate governance risk valued? Evidence from directors' and officers' insurance. Journal of Corporate Finance, 18(2), 349-372.",
    "Bradley, M., & Chen, D. (2011). Corporate governance and the cost of debt: Evidence from director limited liability and indemnification provisions. Journal of Corporate Finance, 17(1), 83-107.",
    "Chalmers, J. M. R., Dann, L. Y., & Harford, J. (2002). Managerial opportunism? Evidence from directors' and officers' insurance purchases. Journal of Finance, 57(2), 609-636.",
    "Core, J. E. (2000). The directors' and officers' insurance premium: An outside assessment of the quality of corporate governance. Journal of Law, Economics, and Organization, 16(2), 449-477.",
    "Donelson, D. C., Hutzler, A. D., & Rhodes, A. (2025). Does antitrust enforcement against interlocking directorates impair corporate governance? Journal of Accounting and Economics.",
    "Dyck, A., Fang, H., Hebert, C., & Xu, J. (2026). Venture fraud and founder-controlled boards. Working paper.",
    "Ewens, M., & Malenko, N. (2024). Board dynamics over the startup lifecycle. Working paper.",
    "Fama, E. F., & Jensen, M. C. (1983). Separation of ownership and control. Journal of Law and Economics, 26(2), 301-325.",
    "Gillan, S. L., & Panasian, C. A. (2015). On lawsuits, corporate governance, and the Dodd-Frank Act. Journal of Corporate Finance, 35, 308-320.",
    "Hermalin, B. E., & Weisbach, M. S. (2003). Boards of directors as an endogenously determined institution: A survey of the economic literature. Economic Policy Review, 9(1), 7-26.",
    "Jensen, M. C., & Meckling, W. H. (1976). Theory of the firm: Managerial behavior, agency costs and ownership structure. Journal of Financial Economics, 3(4), 305-360.",
    "Kaplan, S. N., & Stromberg, P. (2003). Financial contracting theory meets the real world: An empirical analysis of venture capital contracts. Review of Economic Studies, 70(2), 281-315.",
    "Lin, C., Officer, M. S., & Zou, H. (2011). Directors' and officers' liability insurance and acquisition outcomes. Journal of Financial Economics, 102(3), 507-525.",
    "Lin, C., Officer, M. S., Wang, R., & Zou, H. (2013). Directors' and officers' liability insurance and loan spreads. Journal of Financial Economics, 110(1), 37-60.",
    "Lin, C., Officer, M. S., Schmid, T., & Zou, H. (2019). Is skin in the game a game changer? Evidence from mandatory changes of D&O insurance policies. Journal of Accounting and Economics, 68(1).",
    "Nili, Y., et al. (2025). Interlocking directorates in private and public firms. Research handbook. Duke/BU Law.",
    "Packin, N. G., & Alon-Beck, A. (2025). Board observers. University of Illinois Law Review, 2025(5).",
    "Poberejsky, R. (2025). Interlocking directorates, competition, and innovation. Job market paper, Northwestern/Cornerstone Research. SSRN: 4944799.",
    "Raheja, C. G. (2005). Determinants of board size and composition: A theory of corporate boards. Journal of Financial and Quantitative Analysis, 40(2), 283-306.",
    "Shleifer, A., & Vishny, R. W. (1997). A survey of corporate governance. Journal of Finance, 52(2), 737-783.",
]

for ref in refs:
    p = doc.add_paragraph(ref, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(10)

doc.add_paragraph()

# =====================================================================
# DATA SOURCES
# =====================================================================
add_heading_styled("Data Sources and Links", level=1)

sources = [
    ("Ewens & Malenko VC Board Data", "https://github.com/michaelewens/vc_backed_boards"),
    ("Bartlett et al. (2023) Paper", "https://law.stanford.edu/wp-content/uploads/2023/10/2023-09-14_Standardization-and-Innovation-in-Venture-Capital-Contracting.pdf"),
    ("Donelson et al. (2025) — JAE", "https://www.sciencedirect.com/science/article/abs/pii/S0165410125000515"),
    ("Poberejsky (2025) JMP", "https://rpober.github.io/files/JMP_directors.pdf"),
    ("Nili et al. (2025) Interlocks", "https://scholarship.law.duke.edu/faculty_scholarship/4471/"),
    ("Advisen D&O Loss Data", "https://www.advisenltd.com/data/private-d-o-loss-data"),
    ("Woodruff Sawyer — Observer D&O", "https://woodruffsawyer.com/do-notebook/liable-corporate-board-observers/"),
    ("NVCA Model Legal Documents", "https://nvca.org/model-legal-documents/"),
    ("SEC EDGAR Full-Text Search", "https://efts.sec.gov/LATEST/search-index"),
]

for name, url in sources:
    p = doc.add_paragraph(style='List Bullet')
    run_b = p.add_run(name + ": ")
    run_b.bold = True
    run_b.font.size = Pt(10)
    run_u = p.add_run(url)
    run_u.font.size = Pt(10)

# =====================================================================
# SAVE
# =====================================================================
output_path = os.path.join(
    os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
    "..",
    "Four Empirical Settings -- Three-Tier Board Governance.docx"
)
output_path = os.path.normpath(output_path)
doc.save(output_path)
print(f"Saved to: {output_path}")
