"""Edit v3: Add sample attrition tables in Section 3.6."""

import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = docx.Document("C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx")

# Find paragraph 192 (the last paragraph of section 3.6, before Section 4)
# "The attrition pattern also has substantive implications..."
target_idx = None
for i, para in enumerate(doc.paragraphs):
    if "attrition pattern also has substantive implications" in para.text:
        target_idx = i
        break

if target_idx is None:
    # Try alternate
    for i, para in enumerate(doc.paragraphs):
        if "3.6" in para.text and "Attrition" in para.text:
            target_idx = i
            break

print(f"Inserting after paragraph {target_idx}")

body = doc.element.body
insert_after = doc.paragraphs[target_idx]._element

# Helper
def add_after(text, bold=False, italic=False, font_size=12):
    global insert_after
    p = doc.add_paragraph()
    p.style = doc.styles["Normal"]
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    p.paragraph_format.line_spacing = 2.0
    p.paragraph_format.space_after = Pt(6)
    elem = p._element
    body.remove(elem)
    insert_after.addnext(elem)
    insert_after = elem
    return elem

def add_table_after(headers, data, font_size=9):
    global insert_after
    table = doc.add_table(rows=1 + len(data), cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(font_size)
                run.font.name = "Times New Roman"
    for r_idx, row_data in enumerate(data):
        for c_idx, cell_text in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(cell_text)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(font_size)
                    run.font.name = "Times New Roman"
    elem = table._element
    body.remove(elem)
    insert_after.addnext(elem)
    insert_after = elem
    return elem

# =====================================================================
# Table 1: Dataset Overview
# =====================================================================
add_after("Table 3.1: Dataset Overview", bold=True, font_size=11)

add_table_after(
    ["#", "Dataset", "Source", "N (Records)", "N (Unique Cos)", "Key Variables", "Period"],
    [
        ["1", "CIQ Observer Records", "WRDS (ciq_pplintel)", "5,570", "3,058", "Observer titles, person IDs, company links", "As of 2026"],
        ["2", "CIQ Advisory Board", "WRDS (ciq_pplintel)", "180,128", "31,088", "Advisory titles, person IDs", "As of 2026"],
        ["3", "CIQ Directors", "WRDS (ciq_pplintel)", "48,521", "3,046", "All board members at observer cos", "As of 2026"],
        ["4", "CIQ Company Details", "WRDS (ciq_common)", "3,058", "3,058", "Type, status, country, year founded", "As of 2026"],
        ["5", "CIQ Network", "WRDS (ciq_pplintel)", "42,857", "26,079", "All positions held by observers", "As of 2026"],
        ["6", "CIQ Key Dev Events", "WRDS (ciq_keydev)", "61,669", "2,284", "Exec changes, lawsuits, restatements", "2000\u20132026"],
        ["7", "CIQ-CIK Crosswalk", "WRDS (ciq_common)", "1,446", "1,355", "CIQ companyid \u2192 SEC CIK", "\u2014"],
        ["8", "CIQ Deal Amounts", "WRDS (ciq_transactions)", "624", "624", "Total capital raised (USD)", "\u2014"],
        ["9", "EDGAR S-1 Universe", "SEC EDGAR indexes", "28,038", "5,566", "All S-1/S-1A filings", "2017\u20132026"],
        ["10", "EFTS Observer Hits", "SEC EDGAR EFTS API", "1,623", "632", "S-1 files mentioning 'board observer'", "2017\u20132026"],
        ["11", "S-1 Exhibit Coding", "SEC EDGAR exhibits", "320", "204", "Fiduciary language flags", "2017\u20132026"],
        ["12", "Form D Capital", "SEC DERA Form D", "204,604", "204,604", "Total amount sold, industry, revenue range", "2017\u20132025"],
        ["13", "Form D Related Persons", "SEC DERA Form D", "1,711,446", "\u2014", "Director/Officer/Promoter names per filing", "2017\u20132025"],
        ["14", "Ewens-Malenko Panel", "GitHub (public)", "38,369", "7,780", "numVCs, numOut, numExecs by year", "2002\u20132017"],
        ["15", "CRSP Returns", "WRDS (crsp_a_stock)", "\u22483M daily", "\u22482,800", "Daily/monthly stock returns", "2015\u20132026"],
        ["16", "Compustat Annual", "WRDS (comp.funda)", "\u224818,000", "\u22482,600", "Assets, liabilities, sales, net income", "2015\u20132026"],
        ["17", "IBES Consensus", "WRDS (ibes)", "\u2248250,000", "\u22482,100", "Analyst forecasts, dispersion", "2015\u20132026"],
        ["18", "NVCA Model IRA", "NVCA website", "2 docs", "\u2014", "Oct 2023 and Oct 2025 model IRA text", "2023, 2025"],
    ],
    font_size=8
)

add_after(
    "Note: CIQ data is a cross-sectional snapshot as of extraction date (March 2026). "
    "EDGAR and Form D data cover filings from 2017 through 2026 Q1. CRSP/Compustat/IBES "
    "data covers 2015 onward to capture pre-IPO and post-IPO periods.",
    italic=True, font_size=10
)

# =====================================================================
# Table 2: Sample Attrition by Test
# =====================================================================
add_after("Table 3.2: Sample Attrition by Test", bold=True, font_size=11)

add_table_after(
    ["Step", "N", "Attrition", "% Remaining"],
    [
        # Test 1
        ["\u2014 TEST 1: Observer Presence (S-1 Filers) \u2014", "", "", ""],
        ["All S-1 filers (2017\u20132026)", "5,566", "\u2014", "100%"],
        ["  Treatment: EFTS 'board observer' match", "632", "\u2014", "\u2014"],
        ["  Control: No 'board observer' mention", "4,934", "\u2014", "\u2014"],
        ["  \u2192 Matched to CRSP (CIK\u2192GVKEY\u2192PERMNO)", "3,129", "\u22122,437", "56%"],
        ["  \u2192 Have return volatility (\u2265100 trading days)", "2,728", "\u2212401", "49%"],
        ["  \u2192 Have Compustat controls (log assets, leverage)", "1,520", "\u22121,208", "27%"],
        ["    Final: Treatment | Control", "239 | 1,281", "", ""],
        ["    Further: Analyst coverage (IBES match)", "605", "", "11%"],
        ["    Further: IPO underpricing (1st-day return)", "271", "", "5%"],
        # Test 5
        ["", "", "", ""],
        ["\u2014 TEST 5: Full CIQ Private Firms \u2014", "", "", ""],
        ["CIQ observer companies", "3,058", "\u2014", "100%"],
        ["  \u2192 Private Company type only", "2,602", "\u2212456", "85%"],
        ["  \u2192 All controls non-null (board + age + capital)", "2,537", "\u221265", "83%"],
        ["    Attrition: only 2.5%", "", "", ""],
        # Test 3
        ["", "", "", ""],
        ["\u2014 TEST 3: Information Spillover (Network) \u2014", "", "", ""],
        ["CIQ observer persons", "4,915", "\u2014", "100%"],
        ["  \u2192 Have VC/PE affiliation", "2,749", "\u22122,166", "56%"],
        ["  \u2192 Network edges built", "16,670", "\u2014", "\u2014"],
        ["  \u2192 Portfolio co has CRSP PERMNO", "12,885", "\u22123,785", "\u2014"],
        ["  \u2192 Events at observed cos (earnings + exec)", "52,705", "\u2014", "\u2014"],
        ["  \u2192 Event \u00d7 edge pairs", "186,638", "\u2014", "\u2014"],
        ["  \u2192 Daily CARs computed (\u22655 trading days)", "70,218", "\u2212116,420", "\u2014"],
        ["    VC-firm clusters for SEs", "1,353", "", ""],
        # Test 4
        ["", "", "", ""],
        ["\u2014 TEST 4: Pre-Announcement Drift \u2014", "", "", ""],
        ["CRSP-matched S-1 filers", "3,129", "\u2014", "100%"],
        ["  Observer firms in CRSP", "95", "\u2014", "\u2014"],
        ["  Non-observer firms in CRSP", "3,034", "\u2014", "\u2014"],
        ["  \u2192 Events with CRSP coverage", "134,269", "\u2014", "\u2014"],
        ["  \u2192 Sampled + CARs computed", "37,710", "", ""],
        ["    Observer events | Non-observer events", "1,626 | 36,084", "", ""],
    ],
    font_size=8
)

add_after(
    "Note: The critical attrition bottleneck is the CIQ \u2192 CRSP path, which loses 94% of "
    "our sample (3,058 \u2192 184 for Test 2). The CIQ private firm path (Test 5) loses only "
    "2.5%. Test 3 retains statistical power despite attrition through the multiplication "
    "effect (events \u00d7 edges = many observations). For all CRSP-dependent tests, attrition "
    "is driven by private firms lacking SEC CIK identifiers (56% loss), CIKs not linking "
    "to traded securities (further 77% loss), and firms lacking Compustat financial data.",
    italic=True, font_size=10
)

# Save
output = "C:/Users/hjung/Documents/Claude/CorpAcct/Research Proposal v3 -- Three-Tier Board Governance.docx"
doc.save(output)
print(f"Saved updated v3 to: {output}")
