"""Generate paper TOC as Word document."""
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

title = doc.add_heading('Information Permeability of Board Observer Networks', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('Table of Contents \u2014 Draft Outline')
run.font.size = Pt(14)
run.italic = True

doc.add_paragraph()

# Section 1
doc.add_heading('1. Introduction (3-4 pages)', level=1)
for item in [
    'The accountability gap: observers have full information access, zero fiduciary duty',
    '82% of VC entities use board observers (NVCA Q4 2023 CFO Survey)',
    'No empirical business paper has studied observers (Carter, JAE 2025; Ewens & Malenko, 2024, fn. 3)',
    'Preview: same-industry connected firms earn +4.6% abnormal returns before M&A announcements',
    'Two regulatory shocks in opposite directions provide causal identification',
    'Contribution vs interlocking directors literature (private-to-public, pre-event, observer accountability)',
]:
    doc.add_paragraph(item, style='List Bullet')

# Section 2
doc.add_heading('2. Institutional Background (3-4 pages)', level=1)

doc.add_heading('2.1 The Three-Tier Board Architecture', level=2)
doc.add_paragraph('Directors (Tier 1): fiduciary duty, voting rights, personal liability')
doc.add_paragraph('Observers (Tier 2): full information access, no duty, no vote, no liability')
doc.add_paragraph('Advisors (Tier 3): informal, honorary, limited access')

doc.add_heading('2.2 Legal Status of Board Observers', level=2)
doc.add_paragraph('Obasi v. Tibet Pharmaceuticals (3d Cir., 2019): observers are NOT directors')
doc.add_paragraph('No voting rights, no fiduciary duty, tenure not dependent on shareholders')
doc.add_paragraph('Contractual basis: appointed via Investor Rights Agreements (IRAs)')

doc.add_heading('2.3 Regulation Fair Disclosure and the Private Firm Exemption', level=2)
doc.add_paragraph('Reg FD (2000) prohibits selective disclosure by PUBLIC companies')
doc.add_paragraph('Private firms are exempt \u2014 board meetings are outside Reg FD\u2019s reach')
doc.add_paragraph('Observer rights terminate at IPO precisely because of Reg FD')
doc.add_paragraph('The regulatory gap: information crosses from unregulated private to regulated public through the observer\u2019s dual role')

doc.add_heading('2.4 NVCA Model IRA and Observer Provisions', level=2)
doc.add_paragraph('Standard template: observer attends all meetings, receives all materials')
doc.add_paragraph('2020 change: removed \u201cact in a fiduciary manner\u201d clause')
doc.add_paragraph('2025 change: expanded exclusion grounds to \u201ccompetitive harm or competitive disadvantage\u201d')

doc.add_heading('2.5 Clayton Act Section 8 Extension to Observers', level=2)
doc.add_paragraph('Historical scope: directors and officers of competing companies')
doc.add_paragraph('January 2025: DOJ/FTC explicitly extended to board observers')
doc.add_paragraph('Same-industry observer connections now trigger antitrust scrutiny')

# Section 3
doc.add_heading('3. Data and Research Design (4-5 pages)', level=1)

doc.add_heading('3.1 Observer Network', level=2)
doc.add_paragraph('CIQ Professionals: 4,915 observers at 3,058 companies')
doc.add_paragraph('Supplemented with BoardEx (510K crosswalk) and Form 4 (790K crosswalk)')
doc.add_paragraph('Final network: 1,416 observers, 4,747 connected pairs, 2,744 public firms')
doc.add_paragraph('Person-level link: same individual is observer at Firm A and director at Firm B')

doc.add_heading('3.2 Events and Returns', level=2)
doc.add_paragraph('CIQ Key Dev: 400K events, filtered to ~57K (private, not CRSP-listed, no earnings/conferences)')
doc.add_paragraph('Key event types: M&A announcements (with Target/Buyer role), bankruptcy, executive changes')
doc.add_paragraph('CRSP daily returns 2015-2025 (3.3M obs), market-adjusted, winsorized 1/99, exclude price < $5')

doc.add_heading('3.3 Control Group Specification', level=2)
p = doc.add_paragraph('For each event: CARs at ALL portfolio stocks (connected + 10% non-connected sample)')
doc.add_paragraph('CAR = b1(connected) + b2(same_industry) + b3(connected x same_industry) + e')
doc.add_paragraph('Multiple clustering: HC1, event-clustered, stock-clustered')
doc.add_paragraph('Fixed effects: Year FE, Stock FE')

doc.add_heading('3.4 Regulatory Shock Identification', level=2)
doc.add_paragraph('NVCA 2020 (loosen): same_industry x post_2020 interaction')
doc.add_paragraph('Clayton Act 2025 (tighten): same_industry x post_jan2025 interaction')
doc.add_paragraph('Placebos at alternative break years to rule out pre-existing trends')

doc.add_paragraph()
doc.add_paragraph('Table 1: Sample construction and attrition')
doc.add_paragraph('Table 2: Descriptive statistics')

# Section 4
doc.add_heading('4. Results (6-7 pages)', level=1)

doc.add_heading('4.1 Baseline: Connected vs Non-Connected', level=2)
doc.add_paragraph('Exec/Board Changes: connected = +0.28% at CAR[-5,-1] (p=0.010, event-clustered)')
doc.add_paragraph('Survives all non-Stock-FE specifications')
doc.add_paragraph('Establishes the basic information channel')

doc.add_heading('4.2 Event-Type Heterogeneity', level=2)
for item in [
    'M&A Buyer: conn x same = +4.57% at CAR[-30,-1] (p=0.001) \u2014 survives all specs at p<0.01',
    'Bankruptcy: conn x same = +9.45% at CAR[-30,-1] (p=0.043)',
    'M&A Target: conn x same = +1.08% at CAR[-1,0] (p=0.027) \u2014 leaks only at last moment',
    'No effect for: Private Placements, Product/Client, CEO/CFO individually',
    'Information spillover is event-type specific \u2014 only material board decisions',
    'Key: connected alone NOT significant; same_industry alone NOT significant; only the INTERACTION matters',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('4.3 Regulatory Shocks', level=2)
for item in [
    'NVCA 2020: same_ind x post = +3.23% at CAR[-10,-1] (p=0.001). Placebo clean.',
    'Clayton 2025: same_ind x post = -10.19% at CAR[-10,-1] (p<0.001). Placebo clean.',
    'Two shocks, opposite directions, both significant',
    'Four-period trajectory: zero -> positive -> negative -> stays negative',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_paragraph('Table 3: Main results \u2014 all specs as columns, both networks as panels')
doc.add_paragraph('Table 4: By-event-type results')
doc.add_paragraph('Table 5: Shock interaction results')
doc.add_paragraph('Figure 1: Year-by-year same-industry CARs showing regime shift')

# Section 5
doc.add_heading('5. Conclusion (1-2 pages)', level=1)
for item in [
    'Information flows through observer networks; proportional to relevance; responsive to regulation',
    'Framing: information permeability, not insider trading',
    'Regulatory implication: DOJ/FTC\u2019s Clayton Act extension is empirically supported',
    'Limitations: network completeness, mechanism identification, small same-industry N for some tests',
    'Future: public observed companies, calendar-time approach, market model CARs',
]:
    doc.add_paragraph(item, style='List Bullet')

# Appendix
doc.add_page_break()
doc.add_heading('Appendix', level=1)

for item in [
    'A1: Robustness \u2014 SIC3 match, BHARs, connection intensity (dose-response)',
    'A2: Form D as alternative event source',
    'A3: Placebo tests (network shuffle, random event dates)',
    'A4: Abnormal volume analysis',
    'A5: Form 4 insider trading patterns',
    'A6: IRA textual analysis and NVCA template comparison (word-for-word)',
    'A7: Alternative break years for NVCA shock',
    'A8: All 10 CAR windows',
    'A9: Network supplementation details (original vs BoardEx + Form 4)',
]:
    doc.add_paragraph(item, style='List Bullet')

doc.add_paragraph()
doc.add_heading('Figures', level=2)
doc.add_paragraph('Figure 1: Year-by-year same-industry CARs (2015-2025) showing 2020 and 2025 regime shifts')
doc.add_paragraph('Figure 2: Observer network diagram (observer -> VC -> portfolio company)')
doc.add_paragraph('Figure 3: Event-type coefficient comparison (bar chart)')

outpath = 'C:/Users/hjung/Documents/Claude/CorpAcct/Paper TOC -- Board Observer Information Spillover.docx'
doc.save(outpath)
print(f'Saved to {outpath}')
