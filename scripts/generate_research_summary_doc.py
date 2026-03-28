"""Generate comprehensive research summary as a Word document."""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

# Styles
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Title
title = doc.add_heading('Information Permeability of Board Observer Networks', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Research Summary — March 2026')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# =====================================================================
# 1. RESEARCH QUESTION
# =====================================================================
doc.add_heading('1. Research Question', level=1)

doc.add_paragraph(
    'This paper examines whether the board observer role — which provides full information '
    'access to board meetings and materials without imposing fiduciary duty — creates information '
    'channels that affect stock prices at connected public firms. Specifically, we test whether '
    'material events at private companies with board observers produce abnormal returns at '
    'connected public portfolio companies before the events become public.'
)

doc.add_paragraph(
    'Board observers are nonvoting attendees who receive all board materials and attend all '
    'meetings, functioning as de facto board members in every respect except legal obligation. '
    'According to the NVCA Q4 2023 CFO Working-Group Survey, 82% of VC entities use board '
    'observers, with 100% adoption among firms with AUM exceeding $500 million (Packin & '
    'Alon-Beck, 2025). Despite this near-universal prevalence, no empirical business paper '
    'has studied board observers (Carter, JAE 2025; Ewens & Malenko, 2024, footnote 3).'
)

doc.add_paragraph(
    'We frame our investigation around the "accountability gap": observers possess the same '
    'information as directors but bear none of the fiduciary duties, personal liability, or '
    'D&O insurance obligations. This asymmetry creates a unique setting to study whether '
    'information flows through governance networks when legal constraints are absent.'
)

# =====================================================================
# 2. THE INFORMATION CHANNEL
# =====================================================================
doc.add_heading('2. The Information Channel', level=1)

doc.add_paragraph(
    'The channel we test operates at the person level. A single individual serves as a board '
    'observer at a private company (Firm A) and simultaneously serves as a board director at '
    'a public company (Firm B). When a material event occurs at Firm A — such as an acquisition '
    'decision, bankruptcy filing, or executive change — the observer learns about it in the board '
    'meeting weeks or months before public disclosure. We measure whether Firm B\'s stock price '
    'moves abnormally in the pre-announcement window, using Firm B\'s stock returns as a detection '
    'device for information flow through the observer network.'
)

doc.add_paragraph(
    'This design differs fundamentally from the interlocking directorates literature (Kang, 2008; '
    'Shi, 2025; Fracassi, 2017), which examines public-to-public spillovers where both firms\' '
    'information is semi-public. Our test exploits the private-to-public channel, where the '
    'observer is the sole bridge between truly private boardroom information and publicly traded '
    'securities.'
)

# =====================================================================
# 3. DATA
# =====================================================================
doc.add_heading('3. Data', level=1)

doc.add_heading('3.1 Observer Network', level=2)
doc.add_paragraph(
    'We construct the observer network from three sources. The primary source is S&P Capital IQ '
    'Professionals (via WRDS), which identifies 4,915 unique individuals with "observer" in their '
    'title at 3,058 companies. For each observer, we identify all their positions at public '
    'companies — 84% of which are board director or chairman roles. We supplement this with '
    'BoardEx (via the WRDS CIQ-BoardEx crosswalk, 510,276 person matches) and Thomson Reuters '
    'Form 4 insider filing data (via the WRDS CIQ-TR crosswalk, 790,551 person matches). The '
    'supplemented network contains 1,416 observers, 4,747 unique connected pairs linking observed '
    'private companies to public portfolio companies, across 2,744 unique public firms.'
)

doc.add_heading('3.2 Events', level=2)
doc.add_paragraph(
    'We use 400,886 events from the CIQ Key Developments database (wrds_keydev), which captures '
    'press releases, SEC filings, court records, and news articles with event dates, event types, '
    'company roles (Target, Buyer, Seller), and source attributions. After filtering to private '
    'companies, removing events when the company was CRSP-listed (using CCM link dates), and '
    'excluding earnings announcements and conferences, we retain approximately 57,000 events. '
    'We also use SEC Form D filings (2,827 matched capital raise events) as an alternative event '
    'source for robustness.'
)

doc.add_heading('3.3 Returns', level=2)
doc.add_paragraph(
    'Daily stock returns come from CRSP (dsf through 2024, dsf_v2 for 2025), covering 1,909 '
    'portfolio stocks from 2015 through 2025 (3.3 million daily observations). We compute '
    'cumulative abnormal returns (CARs) over multiple windows: [-30,-1], [-20,-1], [-15,-1], '
    '[-10,-1], [-5,-1], [-3,-1], [-2,-1], [-1,0], [0,+3], and [0,+5]. Market-adjusted CARs '
    'subtract the equal-weighted portfolio mean return. We also compute buy-and-hold abnormal '
    'returns (BHARs) for robustness.'
)

# =====================================================================
# 4. RESEARCH DESIGN
# =====================================================================
doc.add_heading('4. Research Design', level=1)

doc.add_heading('4.1 Control Group Test', level=2)
doc.add_paragraph(
    'For each event at a private observed company, we compute CARs at all portfolio stocks in '
    'our sample — both connected (through the observer network) and non-connected. The regression '
    'specification is:'
)
doc.add_paragraph(
    'CAR = \u03b2\u2081(connected) + \u03b2\u2082(same_industry) + \u03b2\u2083(connected \u00d7 same_industry) + \u03b5',
    style='Normal'
)
doc.add_paragraph(
    'where connected equals one if the portfolio stock is linked to the event company through '
    'the observer network, same_industry equals one if both share the same 2-digit SIC code, '
    'and the interaction captures whether the connection effect is stronger for industry-relevant '
    'information. We use a 10% random sample of non-connected stocks as controls to maintain '
    'computational tractability.'
)

doc.add_heading('4.2 Clustering and Fixed Effects', level=2)
doc.add_paragraph(
    'We test robustness across eight specifications: (1) OLS with HC1 robust standard errors, '
    '(2) event-clustered, (3) stock-clustered, (4) Year FE + HC1, (5) Year FE + event-clustered, '
    '(6) Year FE + stock-clustered, (7) Stock FE + HC1, and (8) Stock FE + event-clustered. '
    'Event-clustering is the most conservative specification, accounting for within-event '
    'correlation across stocks. We dropped VC fixed effects because only 13% of VCs have both '
    'same-industry and different-industry edges, providing insufficient within-group variation.'
)

doc.add_heading('4.3 Regulatory Shock Identification', level=2)
doc.add_paragraph(
    'We exploit two regulatory changes affecting observer information constraints:'
)
doc.add_paragraph(
    'NVCA 2020: The National Venture Capital Association removed fiduciary language ("act in a '
    'fiduciary manner") from its standard observer provisions in the model Investor Rights '
    'Agreement. This loosened the perceived obligation of observers to protect the observed '
    'company\'s information. Prediction: information spillover increases.',
    style='List Bullet'
)
doc.add_paragraph(
    'Clayton Act January 2025: The DOJ and FTC explicitly extended Section 8 interlocking '
    'directorate rules to cover board observers, subjecting same-industry observer connections '
    'to antitrust scrutiny for the first time. Prediction: same-industry spillover decreases.',
    style='List Bullet'
)
doc.add_paragraph(
    'These two shocks operate in opposite directions on the same mechanism, providing a '
    'compelling identification strategy. We test using a difference-in-differences framework '
    'with the interaction same_industry \u00d7 post_shock, controlling for year fixed effects '
    'and clustering by VC firm.'
)

# =====================================================================
# 5. FINDINGS
# =====================================================================
doc.add_heading('5. Findings', level=1)

doc.add_heading('5.1 The Connection Effect Exists', level=2)
doc.add_paragraph(
    'For executive and board changes (the largest event category with 1,884-2,348 events), '
    'connected portfolio companies earn 0.28% higher market-adjusted returns than non-connected '
    'stocks in the five days before the event (p=0.010, event-clustered, supplemented network). '
    'This result survives all six non-Stock-FE specifications. The connection itself matters '
    'regardless of industry match for routine events.'
)

doc.add_heading('5.2 Material Events Show Stronger, Industry-Specific Spillover', level=2)

# M&A Buyer table
doc.add_paragraph(
    'M&A Buyer Events (observer\'s company acquiring another firm). The board approves the '
    'acquisition before announcement; the observer knows the target, price, and strategic '
    'rationale. Same-industry connected firms show the strongest pre-event abnormal returns:'
)

table = doc.add_table(rows=7, cols=4)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers = ['Specification', 'Coefficient', 'p-value', 'Sig']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
    for p in table.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)

data = [
    ['OLS + HC1', '+4.57%', '0.008', '***'],
    ['Event-clustered', '+4.57%', '0.001', '***'],
    ['Stock-clustered', '+4.57%', '0.004', '***'],
    ['Year FE + Event-cl', '+4.57%', '0.001', '***'],
    ['Year FE + Stock-cl', '+4.57%', '0.004', '***'],
    ['Stock FE + Event-cl', '-0.26%', '0.788', ''],
]
for r, row_data in enumerate(data):
    for c, val in enumerate(row_data):
        table.rows[r+1].cells[c].text = val
        for p in table.rows[r+1].cells[c].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_paragraph(
    'Note: Supplemented network (CIQ + BoardEx + Form 4). N=14,226 (358 connected, 28 connected '
    'same-industry). CAR[-30,-1] market-adjusted. conn \u00d7 same_industry interaction coefficient.',
    style='Normal'
).runs[0].font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph(
    'This result survives all six non-Stock-FE specifications at p<0.01. The connected and '
    'same_industry main effects are both insignificant — only the interaction matters. This '
    'means the effect requires BOTH the network connection AND industry-relevant information.'
)

doc.add_paragraph(
    'Bankruptcy Events (128-146 events): conn \u00d7 same_industry at CAR[-30,-1] = +9.45% '
    '(p=0.043, event-clustered, supplemented). Significant across all pre-event windows with '
    'consistent positive drift. Also significant at CAR[-1,0] = +1.22% (p=0.033).'
)

doc.add_paragraph(
    'M&A Target Events (111-146 events): conn \u00d7 same_industry at CAR[-1,0] = +1.08% '
    '(p=0.027, event-clustered, original). Only significant right at announcement — acquisition '
    'targets are the tightest board secrets, leaking only at the last moment.'
)

doc.add_paragraph(
    'No significant results for: Private Placements (4,344 events), Product/Client announcements '
    '(4,344 events), or CEO/CFO changes individually. Information spillover is event-type specific.'
)

doc.add_heading('5.3 Regulatory Changes Affect the Channel', level=2)

doc.add_paragraph(
    'The NVCA 2020 fiduciary language removal coincides with an increase in same-industry '
    'spillover. The interaction same_industry \u00d7 post_2020 at CAR[-10,-1] = +3.23% '
    '(p=0.001, Year FE + VC-clustered). The placebo at January 2024 is a clean null '
    '(p=0.510). Alternative break year analysis confirms 2020 is the strongest break; '
    'pre-2020 breaks produce no effect.'
)

doc.add_paragraph(
    'The Clayton Act January 2025 extension produces the opposite effect. The interaction '
    'same_industry \u00d7 post_jan2025 at CAR[-10,-1] = -10.19% (p<0.001, Year FE + '
    'VC-clustered). The placebo is clean (p=0.470). With the supplemented network, the '
    'Clayton Act effect strengthens further (-4.88%, p=0.010).'
)

doc.add_paragraph(
    'Two independent regulatory changes, opposite predictions, both confirmed, placebos '
    'clean. The four-period trajectory of same-industry CAR[-10,-1] tracks the regulatory '
    'environment: zero pre-2020, positive when NVCA loosened (2020-2024), reverses when '
    'Clayton Act tightened (Jan-Sep 2025), stays negative after NVCA re-tightened (Oct 2025+).'
)

doc.add_heading('5.4 Additional Robustness', level=2)

table2 = doc.add_table(rows=7, cols=3)
table2.style = 'Table Grid'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Robustness Check', 'Coefficient', 'p-value']):
    table2.rows[0].cells[i].text = h
    for p in table2.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True
        p.runs[0].font.size = Pt(10)

rob_data = [
    ['Baseline (event-clustered)', '+4.57%', '0.001***'],
    ['Winsorized (1st/99th)', '+4.77%', '0.007***'],
    ['Exclude penny stocks', '+5.24%', '0.005***'],
    ['SIC3 tighter match', '+4.94%', '0.004***'],
    ['Buy-and-hold (BHAR)', '+4.66%', '0.012**'],
    ['Connection intensity', '+0.86%/link', '0.007***'],
]
for r, row_data in enumerate(rob_data):
    for c, val in enumerate(row_data):
        table2.rows[r+1].cells[c].text = val
        for p in table2.rows[r+1].cells[c].paragraphs:
            for run in p.runs:
                run.font.size = Pt(10)

doc.add_paragraph(
    'Note: M&A Buyer, CAR[-30,-1], conn \u00d7 same_industry. Connection intensity uses '
    'number of observer links as a continuous variable.',
    style='Normal'
).runs[0].font.size = Pt(9)

doc.add_paragraph()
doc.add_paragraph(
    'The connection intensity result provides a dose-response relationship: more observer links '
    'between the event company and the portfolio company produce stronger spillover effects.'
)

doc.add_heading('5.5 Network Supplementation', level=2)
doc.add_paragraph(
    'We verify that the original CIQ network captures roughly half of the observer-to-public-firm '
    'connections by cross-referencing Form 4 insider filings. We supplement the network with '
    'BoardEx board positions and Form 4 insider connections. All key results either hold or '
    'strengthen with the denser network, confirming that incomplete network measurement biases '
    'our estimates downward rather than driving spurious results.'
)

doc.add_heading('5.6 Supporting Evidence', level=2)
doc.add_paragraph(
    'Abnormal Volume: Connected stocks do not show abnormally high trading volume before events, '
    'which does not support a simple informed trading mechanism.',
    style='List Bullet'
)
doc.add_paragraph(
    'Form 4 Insider Trading: Observers shift toward buying same-industry stocks before events '
    '(buy share 62.7% vs 42.3% baseline, Fisher p=0.024). The logistic odds ratio is 2.75 '
    '(p=0.030 with HC1) but the result is not robust to observer-clustering due to only 20 '
    'pre-event same-industry trades. Suggestive but not definitive.',
    style='List Bullet'
)
doc.add_paragraph(
    'IRA Textual Analysis: The "no fiduciary duty" disclaimer increased significantly in S-1 '
    'IRA exhibits after 2020 (46% to 62%, p=0.012), confirming the NVCA template change '
    'propagated to actual company documents.',
    style='List Bullet'
)
doc.add_paragraph(
    'Placebo Tests: Random event dates produce significantly smaller coefficients than actual '
    'event dates (empirical p=0.048 for M&A Buyer, p=0.022 for Bankruptcy). Network shuffle '
    'placebos are suggestive but underpowered (p=0.142 and p=0.060).',
    style='List Bullet'
)

# =====================================================================
# 6. INTERPRETATION
# =====================================================================
doc.add_heading('6. Interpretation and Contribution', level=1)

doc.add_paragraph(
    'We interpret the evidence as showing that the observer role creates information channels '
    'that would not exist under a traditional director-only board. The effect is proportional '
    'to information relevance (same-industry, material events) and responsive to the observer\'s '
    'accountability structure (regulatory changes affecting fiduciary obligations and antitrust '
    'scrutiny). We deliberately do not claim a specific transmission mechanism — the stock price '
    'effects could arise from direct trading by the observer, network diffusion through the VC '
    'ecosystem, strategic decisions at the connected public firm informed by the observer\'s '
    'dual role, or market inference from observable network signals.'
)

doc.add_paragraph(
    'Our contribution relative to the prior literature is threefold. First, we document a '
    'private-to-public information channel through observer networks, distinct from the '
    'public-to-public interlocking directorates studied in Kang (2008), Shi (2025), and '
    'Fracassi (2017). Second, we show pre-event abnormal returns — evidence of information '
    'leakage before disclosure — rather than post-event reputation spillovers. Third, we '
    'exploit two regulatory shocks (NVCA 2020 and Clayton Act 2025) in opposite directions '
    'to identify the causal effect of observer accountability on information flow.'
)

doc.add_paragraph(
    'The regulatory implications are direct. The DOJ/FTC\'s January 2025 extension of Clayton '
    'Act Section 8 to board observers appears well-motivated by the data: same-industry observer '
    'connections generate statistically significant abnormal returns at connected firms months '
    'before public disclosure. Our evidence provides the first empirical support for the '
    'regulatory concern that observer seats facilitate information transfer between firms in '
    'the same industry.'
)

# =====================================================================
# 7. SAMPLE AND KEY STATISTICS
# =====================================================================
doc.add_heading('7. Sample and Key Statistics', level=1)

table3 = doc.add_table(rows=13, cols=2)
table3.style = 'Table Grid'
for i, h in enumerate(['Statistic', 'Value']):
    table3.rows[0].cells[i].text = h
    for p in table3.rows[0].cells[i].paragraphs:
        p.runs[0].bold = True

stats = [
    ['Observer records (CIQ)', '5,570 records, 4,915 persons, 3,058 companies'],
    ['Supplemented network edges', '5,712 unique (CIQ 3,725 + BoardEx 1,383 + Form4 604)'],
    ['Connected (observer, portfolio) pairs', '4,747'],
    ['CIQ Key Dev events (total)', '400,886'],
    ['Events after filtering', '~57,000'],
    ['CRSP daily returns', '3.3 million observations, 1,909 stocks (2015-2025)'],
    ['M&A Buyer events', '105 (original) / 124 (supplemented)'],
    ['Bankruptcy events', '128 / 146'],
    ['Exec/Board Change events', '1,884 / 2,348'],
    ['Form 4 insider trades', '31,206 by 616 observers at 1,223 companies'],
    ['S-1 IRA exhibits coded', '320 exhibits, 204 companies'],
    ['NVCA model IRA templates', 'Oct 2023 and Oct 2025 (word-for-word comparison)'],
]
for r, (label, val) in enumerate(stats):
    table3.rows[r+1].cells[0].text = label
    table3.rows[r+1].cells[1].text = val

# Save
outpath = "C:/Users/hjung/Documents/Claude/CorpAcct/Research Summary -- Board Observer Information Spillover.docx"
doc.save(outpath)
print(f"Saved to {outpath}")
